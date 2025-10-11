#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Complete Manuscript Text Extraction
Extracts all text including tables from the final manuscript

Usage: python extract_complete_manuscript.py
"""

from docx import Document
from pathlib import Path

def extract_complete_text(docx_path):
    """Extract all text including tables from manuscript"""
    doc = Document(docx_path)

    all_text = []

    # Extract all paragraphs and tables in order
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            # Find the corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        all_text.append(text)
                    break

        elif element.tag.endswith('tbl'):  # Table
            # Find the corresponding table object
            for table in doc.tables:
                if table._element == element:
                    all_text.append("\n[TABLE START]")
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                            row_text.append(cell_text)
                        if any(row_text):
                            all_text.append(" | ".join(row_text))
                    all_text.append("[TABLE END]\n")
                    break

    return '\n'.join(all_text)

def main():
    """Main execution"""
    manuscript_path = Path(r"C:\jips\docs\manuscript_phase2_complete_20251011_114522.docx")
    output_path = Path(r"C:\jips\docs\manuscript.txt")

    if not manuscript_path.exists():
        print(f"❌ ERROR: Manuscript not found: {manuscript_path}")
        return

    print(f"📄 Extracting text from: {manuscript_path.name}")

    # Extract text
    text = extract_complete_text(manuscript_path)

    # Save to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

    # Statistics
    lines = text.split('\n')
    chars = len(text)
    words = len(text.split())

    print(f"\n✅ Extraction complete!")
    print(f"   Output: {output_path.name}")
    print(f"   Lines: {len(lines):,}")
    print(f"   Characters: {chars:,}")
    print(f"   Words: {words:,}")

    # Quick validation
    print(f"\n🔍 Quick validation:")
    key_checks = [
        ("6.12×", text.count("6.12×")),
        ("15.3%", text.count("15.3%")),
        ("r = 0.987", text.count("r = 0.987")),
        ("κ = 0.260", text.count("κ = 0.260")),
        ("Section 3.1", text.count("3.1")),
        ("Section 6.1", text.count("6.1")),
        ("Appendix C", text.count("Appendix C")),
        ("Appendix E", text.count("Appendix E")),
    ]

    for item, count in key_checks:
        status = "✅" if count > 0 else "❌"
        print(f"   {status} {item}: {count} occurrences")

if __name__ == "__main__":
    main()
