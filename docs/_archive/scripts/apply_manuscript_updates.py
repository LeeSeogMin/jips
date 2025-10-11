#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Automated Manuscript Update Application Script
Applies all Phase 8 updates to the manuscript Word document

Usage: python apply_manuscript_updates.py
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json
from datetime import datetime

class ManuscriptUpdater:
    def __init__(self, manuscript_path, updates_dir):
        self.manuscript_path = Path(manuscript_path)
        self.updates_dir = Path(updates_dir)
        self.doc = Document(str(self.manuscript_path))
        self.backup_path = self.manuscript_path.parent / f"manuscript_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        self.log = []

    def create_backup(self):
        """Create backup of original manuscript"""
        self.doc.save(str(self.backup_path))
        self.log.append(f"✅ Backup created: {self.backup_path.name}")
        print(f"✅ Backup created: {self.backup_path.name}")

    def load_update_file(self, filename):
        """Load update markdown file"""
        filepath = self.updates_dir / filename
        if not filepath.exists():
            self.log.append(f"❌ File not found: {filename}")
            return None

        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        self.log.append(f"✅ Loaded: {filename}")
        return content

    def find_text_in_document(self, search_text, exact=False):
        """Find text in document and return paragraph index"""
        for i, para in enumerate(self.doc.paragraphs):
            if exact:
                if para.text.strip() == search_text.strip():
                    return i
            else:
                if search_text.lower() in para.text.lower():
                    return i
        return None

    def replace_text_global(self, old_text, new_text):
        """Global search and replace in all paragraphs"""
        count = 0
        for para in self.doc.paragraphs:
            if old_text in para.text:
                # Replace inline in runs to preserve formatting
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        count += 1

        # Also replace in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    count += 1

        return count

    def apply_numerical_corrections(self):
        """Phase 1: Apply all numerical corrections"""
        print("\n" + "="*60)
        print("PHASE 1: NUMERICAL CORRECTIONS")
        print("="*60)

        corrections = [
            # Discrimination power
            ("27.3% more accurate", "6.12× better discrimination power (15.3% vs 2.5%)"),
            ("36.5% improvement in discriminative power", "6.12× improvement in discrimination power (15.3% semantic vs 2.5% statistical)"),

            # Correlations
            ("r = 0.88", "r = 0.987"),
            ("r = 0.67", "r = 0.988"),

            # Kappa
            ("Cohen's Kappa (κ = 0.91)", "Fleiss' kappa (κ = 0.260)"),
            ("Cohen's κ = 0.91", "Fleiss' κ = 0.260"),
            ("Cohen's κ", "Fleiss' κ"),
            ("κ = 0.91", "κ = 0.260"),
            ("κ = 0.89", "κ = 0.260"),

            # Inter-topic similarity
            ("average inter-topic similarity of 0.21", "average inter-topic similarity of 0.179"),
            ("shows 0.48", "shows 0.312"),
            ("demonstrates 0.67", "demonstrates 0.358"),

            # Average words (will need manual check for Table 2)
            ("20.24 words", "142.3 words"),
            ("20.04 words", "135.8 words"),
            ("21.48 words", "138.5 words"),
        ]

        total_replacements = 0
        for old, new in corrections:
            count = self.replace_text_global(old, new)
            if count > 0:
                self.log.append(f"  ✅ '{old}' → '{new}' ({count} occurrences)")
                print(f"  ✅ '{old}' → '{new}' ({count} occurrences)")
                total_replacements += count
            else:
                self.log.append(f"  ⚠️  '{old}' not found")
                print(f"  ⚠️  '{old}' not found")

        print(f"\n✅ Phase 1 Complete: {total_replacements} total replacements")
        self.log.append(f"✅ Phase 1 Complete: {total_replacements} total replacements")

    def insert_section_before(self, section_number, new_content, heading):
        """Insert new section before specified section number"""
        # Find the section to insert before
        search_pattern = f"## {section_number}" if section_number.count('.') == 0 else f"### {section_number}"

        target_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.startswith(search_pattern):
                target_idx = i
                break

        if target_idx is None:
            self.log.append(f"  ❌ Could not find section {section_number}")
            print(f"  ❌ Could not find section {section_number}")
            return False

        # Insert new heading
        new_para = self.doc.paragraphs[target_idx].insert_paragraph_before(heading)
        if heading.startswith("### "):
            new_para.style = 'Heading 3'
        elif heading.startswith("## "):
            new_para.style = 'Heading 2'

        # Insert content paragraphs
        for line in new_content.split('\n'):
            if line.strip():
                self.doc.paragraphs[target_idx].insert_paragraph_before(line)

        self.log.append(f"  ✅ Inserted section before {section_number}")
        print(f"  ✅ Inserted section before {section_number}")
        return True

    def replace_section(self, section_number, new_content):
        """Replace entire section content"""
        # Find section start
        search_pattern = f"## {section_number}" if section_number.count('.') == 0 else f"### {section_number}"

        start_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.startswith(search_pattern):
                start_idx = i
                break

        if start_idx is None:
            self.log.append(f"  ❌ Could not find section {section_number}")
            print(f"  ❌ Could not find section {section_number}")
            return False

        # Find section end (next section of same or higher level)
        end_idx = len(self.doc.paragraphs)
        section_level = search_pattern.count('#')

        for i in range(start_idx + 1, len(self.doc.paragraphs)):
            para_text = self.doc.paragraphs[i].text
            if para_text.startswith('#' * section_level + ' '):
                end_idx = i
                break

        # Remove old content (keep heading)
        for i in range(end_idx - 1, start_idx, -1):
            self.doc.paragraphs[i]._element.getparent().remove(self.doc.paragraphs[i]._element)

        # Insert new content after heading
        for line in reversed(new_content.split('\n')):
            if line.strip():
                new_para = self.doc.paragraphs[start_idx].insert_paragraph_after(line)

        self.log.append(f"  ✅ Replaced section {section_number}")
        print(f"  ✅ Replaced section {section_number}")
        return True

    def apply_content_updates(self):
        """Phase 2: Apply content additions section by section"""
        print("\n" + "="*60)
        print("PHASE 2: CONTENT ADDITIONS")
        print("="*60)

        # Note: This is a simplified version. Manual intervention may be needed
        # for complex insertions. The script provides structure for automation.

        print("\n⚠️  Content additions require manual review")
        print("Please use the update files in manuscript_updates/ directory")
        print("Follow the 00_MASTER_UPDATE_GUIDE.md for detailed instructions")

        self.log.append("⚠️  Content additions require manual application")
        self.log.append("See: 00_MASTER_UPDATE_GUIDE.md for instructions")

    def generate_report(self):
        """Generate update application report"""
        report_path = self.manuscript_path.parent / f"update_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

        report = []
        report.append("="*70)
        report.append("MANUSCRIPT UPDATE APPLICATION REPORT")
        report.append("="*70)
        report.append(f"\nDate: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"Original: {self.manuscript_path.name}")
        report.append(f"Backup: {self.backup_path.name}")
        report.append(f"\n" + "-"*70)
        report.append("UPDATE LOG:")
        report.append("-"*70)

        for log_entry in self.log:
            report.append(log_entry)

        report.append("\n" + "-"*70)
        report.append("NEXT STEPS:")
        report.append("-"*70)
        report.append("1. Open manuscript in Word")
        report.append("2. Review all numerical corrections")
        report.append("3. Apply content additions following 00_MASTER_UPDATE_GUIDE.md")
        report.append("4. Run validation checks from the guide")
        report.append("5. Generate final submission version")

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report))

        print(f"\n✅ Report saved: {report_path.name}")
        return report_path

    def save_updated_document(self):
        """Save updated manuscript"""
        output_path = self.manuscript_path.parent / f"manuscript_updated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        self.doc.save(str(output_path))
        self.log.append(f"✅ Updated manuscript saved: {output_path.name}")
        print(f"\n✅ Updated manuscript saved: {output_path.name}")
        return output_path

def main():
    """Main execution function"""
    print("\n" + "="*70)
    print("AUTOMATED MANUSCRIPT UPDATE APPLICATION")
    print("Phase 8: Reviewer Comment Integration")
    print("="*70)

    # Paths
    manuscript_path = Path(r"C:\jips\docs\manuscript.docx")
    updates_dir = Path(r"C:\jips\docs\manuscript_updates")

    # Verify files exist
    if not manuscript_path.exists():
        print(f"❌ ERROR: Manuscript not found: {manuscript_path}")
        return

    if not updates_dir.exists():
        print(f"❌ ERROR: Updates directory not found: {updates_dir}")
        return

    print(f"\n📄 Manuscript: {manuscript_path.name}")
    print(f"📁 Updates: {updates_dir.name}/")

    # Initialize updater
    updater = ManuscriptUpdater(manuscript_path, updates_dir)

    # Step 1: Create backup
    print("\n" + "-"*70)
    print("STEP 1: BACKUP CREATION")
    print("-"*70)
    updater.create_backup()

    # Step 2: Apply numerical corrections (Phase 1)
    updater.apply_numerical_corrections()

    # Step 3: Content additions guidance (Phase 2)
    updater.apply_content_updates()

    # Step 4: Save updated document
    print("\n" + "-"*70)
    print("STEP 3: SAVING UPDATED DOCUMENT")
    print("-"*70)
    output_path = updater.save_updated_document()

    # Step 5: Generate report
    print("\n" + "-"*70)
    print("STEP 4: GENERATING REPORT")
    print("-"*70)
    report_path = updater.generate_report()

    # Final summary
    print("\n" + "="*70)
    print("UPDATE APPLICATION COMPLETE")
    print("="*70)
    print(f"\n📋 Files created:")
    print(f"   - Backup: {updater.backup_path.name}")
    print(f"   - Updated: {output_path.name}")
    print(f"   - Report: {report_path.name}")

    print(f"\n⚠️  IMPORTANT NEXT STEPS:")
    print(f"   1. Open {output_path.name} in Microsoft Word")
    print(f"   2. Review all numerical corrections")
    print(f"   3. Follow 00_MASTER_UPDATE_GUIDE.md for content additions")
    print(f"   4. Run validation checks from the guide")
    print(f"   5. Save final version for submission")

    print(f"\n✅ Phase 1 (Numerical Corrections) complete!")
    print(f"📖 See {report_path.name} for detailed log")

if __name__ == "__main__":
    main()
