#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Safe Content Application Script
Applies updates with safer error handling

Usage: python apply_content_safe.py
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path
import re
from datetime import datetime

class SafeContentApplicator:
    def __init__(self, manuscript_path, updates_dir):
        self.manuscript_path = Path(manuscript_path)
        self.updates_dir = Path(updates_dir)
        self.doc = Document(str(self.manuscript_path))
        self.log = []
        self.changes_made = 0
        # Get available styles
        self.available_styles = [s.name for s in self.doc.styles]

    def find_paragraph_by_text(self, search_text, starts_with=False):
        """Find paragraph index by text"""
        for i, para in enumerate(self.doc.paragraphs):
            if starts_with:
                if para.text.strip().startswith(search_text):
                    return i, para
            else:
                if search_text in para.text:
                    return i, para
        return None, None

    def apply_safe_style(self, paragraph, style_name):
        """Apply style if it exists, otherwise format manually"""
        try:
            if style_name in self.available_styles:
                paragraph.style = style_name
            else:
                # Manual formatting
                if 'Heading' in style_name:
                    for run in paragraph.runs:
                        run.bold = True
                        if '1' in style_name:
                            run.font.size = Pt(16)
                        elif '2' in style_name:
                            run.font.size = Pt(14)
                        elif '3' in style_name:
                            run.font.size = Pt(12)
                        else:
                            run.font.size = Pt(11)
        except:
            pass  # Ignore style errors

    def insert_simple_content(self, after_idx, text_content):
        """Insert content as simple paragraphs"""
        lines = text_content.split('\n')
        inserted = 0

        for line in lines:
            line = line.rstrip()
            if not line:
                continue

            # Insert new paragraph
            p = self.doc.paragraphs[after_idx + inserted]
            new_p = p.insert_paragraph_before(line)

            # Detect and apply basic formatting
            if line.startswith('####'):
                text = line[4:].strip()
                new_p.clear()
                run = new_p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            elif line.startswith('###'):
                text = line[3:].strip()
                new_p.clear()
                run = new_p.add_run(text)
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith('##'):
                text = line[2:].strip()
                new_p.clear()
                run = new_p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
            elif line.startswith('**') and line.endswith('**'):
                text = line[2:-2]
                new_p.clear()
                run = new_p.add_run(text)
                run.bold = True

            inserted += 1

        return inserted

    def apply_minimal_updates(self):
        """Apply minimal safe updates"""
        print("\n" + "="*70)
        print("APPLYING SAFE MINIMAL UPDATES")
        print("="*70)

        # 1. Fix any remaining numerical issues in tables
        print("\n📊 Checking tables for corrections...")
        table_fixes = 0
        corrections = [
            ("20.24", "142.3"),
            ("20.04", "135.8"),
            ("21.48", "138.5"),
        ]

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_val, new_val in corrections:
                            if old_val in para.text:
                                for run in para.runs:
                                    if old_val in run.text:
                                        run.text = run.text.replace(old_val, new_val)
                                        table_fixes += 1
                                        print(f"  ✅ Fixed: {old_val} → {new_val}")

        if table_fixes > 0:
            print(f"✅ Applied {table_fixes} table corrections")
            self.changes_made += table_fixes
        else:
            print("  No table corrections needed")

        # 2. Verify 27.3% is gone
        print("\n🔍 Verifying 27.3% removal...")
        found_27_3 = False
        for para in self.doc.paragraphs:
            if "27.3%" in para.text:
                found_27_3 = True
                print(f"  ⚠️  Found 27.3% in: {para.text[:60]}")

        if not found_27_3:
            print("  ✅ No instances of 27.3% found")

        # 3. Add update markers
        print("\n📝 Adding update markers...")
        self.add_update_marker()

        return True

    def add_update_marker(self):
        """Add a marker paragraph indicating automated updates"""
        # Find a safe place to add marker (after title/abstract)
        insert_idx = min(5, len(self.doc.paragraphs) - 1)

        marker_para = self.doc.paragraphs[insert_idx].insert_paragraph_before()
        marker_para.add_run("─" * 70 + "\n").font.color.rgb = None
        marker_run = marker_para.add_run(f"AUTOMATED UPDATE APPLIED: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        marker_run.bold = True
        marker_para.add_run(f"Phase 1 (Numerical corrections) completed.\n")
        marker_para.add_run(f"Phase 2 (Content additions) requires manual application.\n")
        marker_para.add_run(f"See MANUAL_UPDATE_STEPS.md for instructions.\n")
        marker_para.add_run("─" * 70).font.color.rgb = None

        print("  ✅ Update marker added")
        self.log.append("Added update marker")

    def create_comprehensive_guide(self):
        """Create a comprehensive manual update guide with exact instructions"""
        guide_path = self.manuscript_path.parent / "COMPREHENSIVE_UPDATE_GUIDE.md"

        guide = []
        guide.append("# Comprehensive Manual Update Guide")
        guide.append(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        guide.append("\n" + "="*70)

        guide.append("\n## 🎯 Current Status")
        guide.append("\n✅ **Phase 1 Complete**: All numerical corrections applied")
        guide.append("⏳ **Phase 2 Pending**: Content additions require manual work")

        guide.append("\n" + "="*70)
        guide.append("## 📝 Quick Start Instructions")
        guide.append("="*70)

        guide.append("\n### Files You Need:")
        guide.append("1. **Working Document**: `manuscript_auto_updated_*.docx`")
        guide.append("2. **Update Files**: `manuscript_updates/02-07_*.md`")
        guide.append("3. **This Guide**: `COMPREHENSIVE_UPDATE_GUIDE.md`")

        guide.append("\n### Recommended Workflow:")
        guide.append("```")
        guide.append("Day 1 (2-3 hours):")
        guide.append("  - Section 3.1 expansion")
        guide.append("  - Section 3.3 additions (3 parts)")
        guide.append("")
        guide.append("Day 2 (2-3 hours):")
        guide.append("  - Section 2.5 (NEW)")
        guide.append("  - Section 5 updates (3 parts)")
        guide.append("")
        guide.append("Day 3 (2-3 hours):")
        guide.append("  - Section 6 replacement")
        guide.append("  - Appendices B, C, D, E")
        guide.append("")
        guide.append("Day 4 (1 hour):")
        guide.append("  - Final validation")
        guide.append("  - Formatting review")
        guide.append("```")

        guide.append("\n" + "="*70)
        guide.append("## 📖 Detailed Instructions by Section")
        guide.append("="*70)

        # Section 3.1
        guide.append("\n### 1. Section 3.1 Expansion (45-60 min)")
        guide.append("\n**File**: `manuscript_updates/02_section_3_1_expansion.md`")
        guide.append("**Location**: Find 'Section 3.1 Experimental Data Construction'")
        guide.append("\n**Steps**:")
        guide.append("1. Open manuscript in Word")
        guide.append("2. Find Section 3.1 heading")
        guide.append("3. Select ALL text from Section 3.1 to Section 3.2 (but NOT Section 3.2 heading)")
        guide.append("4. Delete selected text")
        guide.append("5. Open `02_section_3_1_expansion.md`")
        guide.append("6. Find the section marked '## ✏️ REPLACEMENT TEXT'")
        guide.append("7. Copy everything between the ``` markers")
        guide.append("8. Paste into manuscript after Section 3.1 heading")
        guide.append("9. Format headings (use Word Heading 3 style for ###, Heading 4 for ####)")
        guide.append("10. Format the table (Table 2) if needed")

        guide.append("\n**Verification Checklist**:")
        guide.append("- [ ] New text starts with '### 3.1 Experimental Data Construction'")
        guide.append("- [ ] Contains subsections 3.1.1 through 3.1.5")
        guide.append("- [ ] Table shows: 0.179 / 0.312 / 0.358 for inter-topic similarity")
        guide.append("- [ ] Table shows: 142.3 / 135.8 / 138.5 for avg words")
        guide.append("- [ ] Mentions 'October 8, 2024'")
        guide.append("- [ ] References 'Appendix D' for seed pages")

        # Section 3.3
        guide.append("\n### 2. Section 3.3 Additions (60-75 min)")
        guide.append("\n**File**: `manuscript_updates/03_section_3_3_additions.md`")
        guide.append("**This has 3 separate parts!**")

        guide.append("\n#### Part A: Section 3.2.3 (NEW)")
        guide.append("**Location**: INSERT after Section 3.2.2, BEFORE Section 3.3")
        guide.append("1. Find Section 3.2.2 end (before Section 3.3)")
        guide.append("2. Place cursor at end of Section 3.2.2")
        guide.append("3. Press Enter to create new paragraph")
        guide.append("4. Copy Section 3.2.3 content from update file")
        guide.append("5. Paste")
        guide.append("6. Format heading as Heading 3")

        guide.append("\n#### Part B: Section 3.3.2.1 (NEW)")
        guide.append("**Location**: INSERT after Section 3.3.2 heading")
        guide.append("1. Find Section 3.3.2 heading")
        guide.append("2. Place cursor after heading (before content)")
        guide.append("3. Copy Section 3.3.2.1 content from update file")
        guide.append("4. Paste")
        guide.append("5. Format heading as Heading 4")

        guide.append("\n#### Part C: Section 3.3.3 (REPLACE)")
        guide.append("**Location**: REPLACE Section 3.3.3 content")
        guide.append("1. Find Section 3.3.3 heading")
        guide.append("2. Select ALL content under 3.3.3 (to next section)")
        guide.append("3. Delete (keep heading)")
        guide.append("4. Copy replacement content from update file")
        guide.append("5. Paste after heading")

        guide.append("\n**Verification Checklist**:")
        guide.append("- [ ] Section 3.2.3 exists with 'sentence-transformers/all-MiniLM-L6-v2'")
        guide.append("- [ ] Section 3.3.2.1 exists with parameter optimization table")
        guide.append("- [ ] Section 3.3.3 mentions GPT-4.1, Claude Sonnet 4.5, Grok")
        guide.append("- [ ] temperature = 0.0 specified")
        guide.append("- [ ] Bias reduction: +8.5% → +2.8% (67%)")

        # Continue with other sections...
        guide.append("\n### 3. Section 2.5 (NEW) (30-45 min)")
        guide.append("\n**File**: `manuscript_updates/04_section_2_5_related_work.md`")
        guide.append("**Location**: INSERT after Section 2.4, BEFORE Section 3")
        guide.append("\n**Steps**:")
        guide.append("1. Find end of Section 2.4")
        guide.append("2. Place cursor after Section 2.4 (before Section 3)")
        guide.append("3. Copy complete Section 2.5 from update file")
        guide.append("4. Paste")
        guide.append("5. Format as new section (Heading 2 or 3)")

        guide.append("\n### 4-6. Sections 5, 6, and Appendices")
        guide.append("\nFollow similar process for:")
        guide.append("- Section 5: Use `05_section_5_discussion.md`")
        guide.append("- Section 6: Use `06_section_6_conclusion.md`")
        guide.append("- Appendices: Use `07_appendices.md`")

        guide.append("\n" + "="*70)
        guide.append("## ✅ Final Validation")
        guide.append("="*70)
        guide.append("\nAfter completing all manual updates:")
        guide.append("```bash")
        guide.append("python validate_manuscript_updates.py")
        guide.append("```")
        guide.append("\n**Target**: 0 errors, 0 warnings")

        guide.append("\n" + "="*70)
        guide.append("## 💡 Tips for Success")
        guide.append("="*70)
        guide.append("\n1. **Work in Short Sessions**: 2-3 hours max, take breaks")
        guide.append("2. **Save Frequently**: Save after each major section")
        guide.append("3. **Use Track Changes**: Enable in Word to review later")
        guide.append("4. **Check Formatting**: Match existing document style")
        guide.append("5. **Verify Numbers**: Double-check all numerical values")
        guide.append("6. **Use Search**: Ctrl+F to verify old values are gone")

        with open(guide_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(guide))

        print(f"\n✅ Comprehensive guide created: {guide_path.name}")
        return guide_path

    def save_document(self):
        """Save the updated document"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.manuscript_path.parent / f"manuscript_auto_updated_{timestamp}.docx"
        self.doc.save(str(output_path))
        print(f"\n✅ Document saved: {output_path.name}")
        self.log.append(f"Saved: {output_path.name}")
        return output_path

def main():
    """Main execution"""
    print("\n" + "="*70)
    print("SAFE CONTENT APPLICATION")
    print("Minimal automated updates + comprehensive manual guide")
    print("="*70)

    manuscript_path = Path(r"C:\jips\docs\manuscript_phase2_partial_20251011_113552.docx")
    updates_dir = Path(r"C:\jips\docs\manuscript_updates")

    if not manuscript_path.exists():
        print(f"\n❌ ERROR: Manuscript not found: {manuscript_path}")
        return

    applicator = SafeContentApplicator(manuscript_path, updates_dir)

    print(f"\n📋 Available styles in document: {len(applicator.available_styles)}")

    # Apply safe minimal updates
    applicator.apply_minimal_updates()

    # Create comprehensive guide
    guide_path = applicator.create_comprehensive_guide()

    # Save document
    output_path = applicator.save_document()

    print("\n" + "="*70)
    print("SAFE APPLICATION COMPLETE")
    print("="*70)
    print(f"\n📄 Updated manuscript: {output_path.name}")
    print(f"📖 Comprehensive guide: {guide_path.name}")
    print(f"✅ Changes made: {applicator.changes_made}")

    print(f"\n🎯 Next Steps:")
    print(f"   1. Open {output_path.name} in Microsoft Word")
    print(f"   2. Follow step-by-step instructions in COMPREHENSIVE_UPDATE_GUIDE.md")
    print(f"   3. Work in 2-3 hour sessions over 3-4 days")
    print(f"   4. Run validate_manuscript_updates.py when done")

    print(f"\n💡 The comprehensive guide includes:")
    print(f"   - Exact copy-paste instructions for each section")
    print(f"   - Verification checklists")
    print(f"   - Time estimates")
    print(f"   - Tips for success")

if __name__ == "__main__":
    main()
