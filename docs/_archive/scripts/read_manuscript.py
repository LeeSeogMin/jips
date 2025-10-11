"""
Read and extract text from manuscript.docx
"""

from docx import Document
import sys

def read_docx(filepath):
    """Extract text from docx file"""
    try:
        doc = Document(filepath)

        full_text = []

        # Extract title and headings
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

    except Exception as e:
        print(f"Error reading docx: {e}", file=sys.stderr)
        return None

if __name__ == "__main__":
    filepath = "docs/manuscript.docx"
    text = read_docx(filepath)

    if text:
        # Print first 5000 characters for preview
        print("="*80)
        print("MANUSCRIPT PREVIEW (first 5000 characters)")
        print("="*80)
        print(text[:5000])
        print("\n...")
        print("="*80)
        print(f"Total length: {len(text)} characters")
        print(f"Total lines: {text.count(chr(10))} lines")

        # Save to text file for analysis
        with open("docs/manuscript_extracted.txt", "w", encoding="utf-8") as f:
            f.write(text)
        print("Full text saved to: docs/manuscript_extracted.txt")
    else:
        print("Failed to extract text from manuscript")
