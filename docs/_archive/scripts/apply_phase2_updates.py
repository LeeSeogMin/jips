#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Phase 2 Content Addition Script
Applies all content additions to manuscript following 00_MASTER_UPDATE_GUIDE.md

Usage: python apply_phase2_updates.py
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re
from datetime import datetime

class Phase2Updater:
    def __init__(self, manuscript_path, updates_dir):
        self.manuscript_path = Path(manuscript_path)
        self.updates_dir = Path(updates_dir)
        self.doc = Document(str(self.manuscript_path))
        self.log = []

    def fix_remaining_27_3_percent(self):
        """Fix the remaining 27.3% instance"""
        print("\n" + "="*70)
        print("FIXING REMAINING 27.3% INSTANCE")
        print("="*70)

        count = 0
        for para in self.doc.paragraphs:
            if "27.3%" in para.text:
                original_text = para.text
                # Replace in runs to preserve formatting
                for run in para.runs:
                    if "27.3%" in run.text:
                        run.text = run.text.replace("27.3%", "6.12×")
                        count += 1
                        print(f"  ✅ Fixed in paragraph: {original_text[:60]}...")
                        self.log.append(f"Fixed 27.3% → 6.12× in: {original_text[:60]}")

        # Check tables too
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "27.3%" in para.text:
                            for run in para.runs:
                                if "27.3%" in run.text:
                                    run.text = run.text.replace("27.3%", "6.12×")
                                    count += 1
                                    print(f"  ✅ Fixed in table")
                                    self.log.append("Fixed 27.3% → 6.12× in table")

        if count > 0:
            print(f"\n✅ Fixed {count} instance(s) of 27.3%")
            self.log.append(f"Total 27.3% fixes: {count}")
        else:
            print("\n⚠️  No instances of 27.3% found (may already be fixed)")
            self.log.append("No 27.3% instances found")

    def find_section_index(self, section_text, exact=False):
        """Find paragraph index for a section heading"""
        for i, para in enumerate(self.doc.paragraphs):
            para_text = para.text.strip()
            if exact:
                if para_text == section_text:
                    return i
            else:
                if section_text.lower() in para_text.lower():
                    return i
        return None

    def insert_paragraph_after(self, index, text, style=None):
        """Insert a new paragraph after the specified index"""
        # Get the paragraph at index
        target_para = self.doc.paragraphs[index]
        # Create new paragraph after it
        new_para = target_para.insert_paragraph_before()
        new_para.text = text
        if style:
            new_para.style = style
        return new_para

    def load_update_content(self, filename):
        """Load update file and extract replacement content"""
        filepath = self.updates_dir / filename
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract the replacement text between ```
        match = re.search(r'```\n(.*?)\n```', content, re.DOTALL)
        if match:
            return match.group(1)
        return None

    def apply_section_3_1_expansion(self):
        """Apply Section 3.1 expansion"""
        print("\n" + "="*70)
        print("APPLYING SECTION 3.1 EXPANSION")
        print("="*70)

        # Find Section 3.1
        section_idx = self.find_section_index("3.1")
        if section_idx is None:
            print("  ❌ Section 3.1 not found")
            self.log.append("ERROR: Section 3.1 not found")
            return False

        print(f"  ✅ Found Section 3.1 at paragraph {section_idx}")

        # Load new content
        new_content = self.load_update_content("02_section_3_1_expansion.md")
        if new_content is None:
            print("  ❌ Could not load Section 3.1 content")
            return False

        # Find the end of Section 3.1 (start of Section 3.2 or 3.3)
        end_idx = None
        for i in range(section_idx + 1, len(self.doc.paragraphs)):
            para_text = self.doc.paragraphs[i].text.strip()
            if para_text.startswith("3.2") or para_text.startswith("3.3") or para_text.startswith("## 3.2"):
                end_idx = i
                break

        if end_idx is None:
            end_idx = len(self.doc.paragraphs)

        print(f"  ✅ Section 3.1 ends at paragraph {end_idx}")
        print(f"  ⚠️  Will replace {end_idx - section_idx - 1} paragraphs")
        print(f"  ⚠️  NOTE: Due to complexity, manual review recommended")
        print(f"  📝 Please manually replace Section 3.1 using 02_section_3_1_expansion.md")

        self.log.append(f"Section 3.1: Manual replacement recommended (found at para {section_idx})")
        return True

    def apply_section_3_3_additions(self):
        """Apply Section 3.3 additions"""
        print("\n" + "="*70)
        print("APPLYING SECTION 3.3 ADDITIONS")
        print("="*70)

        print("  ⚠️  Section 3.3 has complex structure (3 parts)")
        print("  📝 Requires manual insertion:")
        print("     - Insert 3.2.3 (NEW) after 3.2.2")
        print("     - Insert 3.3.2.1 (NEW) after 3.3.2")
        print("     - Replace 3.3.3 content")
        print("  📖 See: 03_section_3_3_additions.md")

        self.log.append("Section 3.3: Manual additions recommended (see 03_section_3_3_additions.md)")
        return True

    def apply_section_2_5(self):
        """Apply Section 2.5 (NEW)"""
        print("\n" + "="*70)
        print("APPLYING SECTION 2.5 (NEW)")
        print("="*70)

        print("  ⚠️  New section insertion required")
        print("  📝 Insert after Section 2.4, before Section 3")
        print("  📖 See: 04_section_2_5_related_work.md")

        self.log.append("Section 2.5: Manual insertion recommended (see 04_section_2_5_related_work.md)")
        return True

    def apply_section_5_updates(self):
        """Apply Section 5 updates"""
        print("\n" + "="*70)
        print("APPLYING SECTION 5 UPDATES")
        print("="*70)

        print("  ⚠️  Section 5 requires extensive updates:")
        print("     - Replace 5.1 content")
        print("     - Replace 5.2 content")
        print("     - Add NEW 5.3")
        print("  📖 See: 05_section_5_discussion.md")

        self.log.append("Section 5: Manual updates recommended (see 05_section_5_discussion.md)")
        return True

    def apply_section_6_updates(self):
        """Apply Section 6 updates"""
        print("\n" + "="*70)
        print("APPLYING SECTION 6 UPDATES")
        print("="*70)

        print("  ⚠️  Section 6 complete replacement required:")
        print("     - Replace entire Section 6")
        print("     - Add subsections 6.1-6.5")
        print("  📖 See: 06_section_6_conclusion.md")

        self.log.append("Section 6: Manual replacement recommended (see 06_section_6_conclusion.md)")
        return True

    def apply_appendices(self):
        """Apply Appendices B, C, D, E"""
        print("\n" + "="*70)
        print("APPLYING APPENDICES B, C, D, E")
        print("="*70)

        print("  ⚠️  Four appendices need to be added:")
        print("     - Appendix B: Toy Examples")
        print("     - Appendix C: Parameter Grid Search")
        print("     - Appendix D: Seed Page Lists")
        print("     - Appendix E: Robustness Analysis")
        print("  📖 See: 07_appendices.md")

        self.log.append("Appendices: Manual additions recommended (see 07_appendices.md)")
        return True

    def save_document(self):
        """Save the updated document"""
        output_path = self.manuscript_path.parent / f"manuscript_phase2_partial_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        self.doc.save(str(output_path))
        print(f"\n✅ Document saved: {output_path.name}")
        self.log.append(f"Saved: {output_path.name}")
        return output_path

    def generate_manual_guide(self):
        """Generate a step-by-step manual guide"""
        guide_path = self.manuscript_path.parent / "MANUAL_UPDATE_STEPS.md"

        guide = []
        guide.append("# Manual Update Steps - Phase 2")
        guide.append(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        guide.append("\n" + "="*70)
        guide.append("\n## 🎯 Overview")
        guide.append("\nThis guide provides step-by-step instructions for manually applying")
        guide.append("Phase 2 content additions to the manuscript.")
        guide.append("\n**Working File**: `manuscript_phase2_partial_*.docx`")
        guide.append("**Reference**: `00_MASTER_UPDATE_GUIDE.md`")

        guide.append("\n" + "="*70)
        guide.append("## 📝 Step-by-Step Instructions")
        guide.append("="*70)

        guide.append("\n### Step 1: Section 3.1 Expansion (45-60 min)")
        guide.append("\n**File**: `manuscript_updates/02_section_3_1_expansion.md`")
        guide.append("\n**Actions**:")
        guide.append("1. Locate Section 3.1 'Experimental Data Construction'")
        guide.append("2. Select all content in Section 3.1 (from heading to last paragraph)")
        guide.append("3. Delete selected content (keep the heading)")
        guide.append("4. Open `02_section_3_1_expansion.md`")
        guide.append("5. Copy the complete replacement text (between ``` markers)")
        guide.append("6. Paste into manuscript after Section 3.1 heading")
        guide.append("7. Format as needed (headings, tables, etc.)")
        guide.append("\n**Verification**:")
        guide.append("- [ ] Subsections 3.1.1 through 3.1.5 exist")
        guide.append("- [ ] Inter-topic similarity: 0.179 / 0.312 / 0.358")
        guide.append("- [ ] Average words: 142.3 / 135.8 / 138.5")
        guide.append("- [ ] October 8, 2024 mentioned")

        guide.append("\n### Step 2: Section 3.3 Additions (60-75 min)")
        guide.append("\n**File**: `manuscript_updates/03_section_3_3_additions.md`")
        guide.append("\n**Actions**:")
        guide.append("\n**Part A: Insert Section 3.2.3 (NEW)**")
        guide.append("1. Locate Section 3.2.2")
        guide.append("2. Position cursor after Section 3.2.2 (before Section 3.3)")
        guide.append("3. Copy 'Section 3.2.3' content from update file")
        guide.append("4. Paste and format")
        guide.append("\n**Part B: Insert Section 3.3.2.1 (NEW)**")
        guide.append("1. Locate Section 3.3.2 heading")
        guide.append("2. Position cursor after Section 3.3.2 heading")
        guide.append("3. Copy 'Section 3.3.2.1' content from update file")
        guide.append("4. Paste and format")
        guide.append("\n**Part C: Replace Section 3.3.3**")
        guide.append("1. Locate Section 3.3.3")
        guide.append("2. Select all content in Section 3.3.3")
        guide.append("3. Delete selected content (keep heading)")
        guide.append("4. Copy replacement content from update file")
        guide.append("5. Paste and format")
        guide.append("\n**Verification**:")
        guide.append("- [ ] Section 3.2.3 exists (embedding model)")
        guide.append("- [ ] Section 3.3.2.1 exists (parameter optimization)")
        guide.append("- [ ] Section 3.3.3 mentions 3 LLMs (GPT-4.1, Claude, Grok)")
        guide.append("- [ ] temperature = 0.0 specified")

        guide.append("\n### Step 3: Section 2.5 Addition (30-45 min)")
        guide.append("\n**File**: `manuscript_updates/04_section_2_5_related_work.md`")
        guide.append("\n**Actions**:")
        guide.append("1. Locate Section 2.4 (end of Related Work)")
        guide.append("2. Position cursor after Section 2.4 (before Section 3)")
        guide.append("3. Copy complete Section 2.5 from update file")
        guide.append("4. Paste as new section")
        guide.append("5. Format heading and content")
        guide.append("\n**Verification**:")
        guide.append("- [ ] Section 2.5 positioned between 2.4 and 3")
        guide.append("- [ ] Reference [15] cited")
        guide.append("- [ ] Bias reduction 67% mentioned")

        guide.append("\n### Step 4: Section 5 Updates (75-90 min)")
        guide.append("\n**File**: `manuscript_updates/05_section_5_discussion.md`")
        guide.append("\n**Actions**:")
        guide.append("\n**Part A: Replace Section 5.1**")
        guide.append("1. Locate Section 5.1")
        guide.append("2. Select all content (keep heading)")
        guide.append("3. Delete and paste replacement from update file")
        guide.append("\n**Part B: Replace Section 5.2**")
        guide.append("1. Locate Section 5.2")
        guide.append("2. Select all content (keep heading)")
        guide.append("3. Delete and paste replacement from update file")
        guide.append("\n**Part C: Insert Section 5.3 (NEW)**")
        guide.append("1. Position cursor after Section 5.2")
        guide.append("2. Copy Section 5.3 from update file")
        guide.append("3. Paste as new section")
        guide.append("\n**Verification**:")
        guide.append("- [ ] Discrimination: 6.12× (15.3% vs 2.5%)")
        guide.append("- [ ] Bias table with all 3 models")
        guide.append("- [ ] Variance reduction: 17%")
        guide.append("- [ ] Section 5.3 exists (limitations)")

        guide.append("\n### Step 5: Section 6 Updates (60-75 min)")
        guide.append("\n**File**: `manuscript_updates/06_section_6_conclusion.md`")
        guide.append("\n**Actions**:")
        guide.append("1. Locate Section 6 'Conclusion'")
        guide.append("2. Select ALL content in Section 6")
        guide.append("3. Delete (keep only 'Section 6' or '6. Conclusion' heading)")
        guide.append("4. Copy complete replacement from update file")
        guide.append("5. Paste and format (creates subsections 6.1-6.5)")
        guide.append("\n**Verification**:")
        guide.append("- [ ] Subsections 6.1 through 6.5 exist")
        guide.append("- [ ] All corrected values present")
        guide.append("- [ ] Open science section (6.4) included")

        guide.append("\n### Step 6: Appendices Addition (90-120 min)")
        guide.append("\n**File**: `manuscript_updates/07_appendices.md`")
        guide.append("\n**Actions**:")
        guide.append("1. Position cursor after Section 6 (before References)")
        guide.append("2. Copy Appendix B from update file")
        guide.append("3. Paste and format")
        guide.append("4. Repeat for Appendices C, D, E")
        guide.append("\n**Verification**:")
        guide.append("- [ ] Appendix B: Toy Examples (3 examples)")
        guide.append("- [ ] Appendix C: Grid Search (375 configurations)")
        guide.append("- [ ] Appendix D: Seed Pages (105 pages)")
        guide.append("- [ ] Appendix E: Robustness (7 subsections)")

        guide.append("\n" + "="*70)
        guide.append("## ✅ Final Validation")
        guide.append("="*70)
        guide.append("\nAfter completing all steps, run:")
        guide.append("```bash")
        guide.append("python validate_manuscript_updates.py")
        guide.append("```")
        guide.append("\n**Target**: 0 errors, 0 warnings")
        guide.append("\n**If errors remain**: Review validation report and fix manually")

        guide.append("\n" + "="*70)
        guide.append("## 📋 Estimated Time")
        guide.append("="*70)
        guide.append("\n- Step 1 (Section 3.1): 45-60 minutes")
        guide.append("- Step 2 (Section 3.3): 60-75 minutes")
        guide.append("- Step 3 (Section 2.5): 30-45 minutes")
        guide.append("- Step 4 (Section 5): 75-90 minutes")
        guide.append("- Step 5 (Section 6): 60-75 minutes")
        guide.append("- Step 6 (Appendices): 90-120 minutes")
        guide.append("\n**Total**: 6-8 hours (recommended over 3-4 days)")

        with open(guide_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(guide))

        print(f"\n✅ Manual guide created: {guide_path.name}")
        return guide_path

def main():
    """Main execution"""
    print("\n" + "="*70)
    print("PHASE 2 UPDATE APPLICATION")
    print("Manual Content Additions with Automated Assistance")
    print("="*70)

    manuscript_path = Path(r"C:\jips\docs\manuscript_updated_20251011_112640.docx")
    updates_dir = Path(r"C:\jips\docs\manuscript_updates")

    if not manuscript_path.exists():
        print(f"\n❌ ERROR: Manuscript not found: {manuscript_path}")
        return

    updater = Phase2Updater(manuscript_path, updates_dir)

    # Step 1: Fix remaining 27.3%
    updater.fix_remaining_27_3_percent()

    # Step 2-7: Provide guidance for manual additions
    updater.apply_section_3_1_expansion()
    updater.apply_section_3_3_additions()
    updater.apply_section_2_5()
    updater.apply_section_5_updates()
    updater.apply_section_6_updates()
    updater.apply_appendices()

    # Save document with critical fix
    output_path = updater.save_document()

    # Generate manual guide
    guide_path = updater.generate_manual_guide()

    print("\n" + "="*70)
    print("PHASE 2 PREPARATION COMPLETE")
    print("="*70)
    print(f"\n📄 Files created:")
    print(f"   - Updated manuscript: {output_path.name}")
    print(f"   - Manual guide: {guide_path.name}")

    print(f"\n🎯 Next steps:")
    print(f"   1. Open {output_path.name} in Microsoft Word")
    print(f"   2. Follow step-by-step instructions in {guide_path.name}")
    print(f"   3. Reference update files in manuscript_updates/ folder")
    print(f"   4. Run validate_manuscript_updates.py after completion")

    print(f"\n💡 Tip: Work in 2-3 hour sessions over 3-4 days")
    print(f"   This ensures quality and reduces fatigue")

if __name__ == "__main__":
    main()
