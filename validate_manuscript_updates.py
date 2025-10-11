#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Manuscript Update Validation Script
Validates all Phase 8 updates have been correctly applied

Usage: python validate_manuscript_updates.py
"""

from docx import Document
from pathlib import Path
import re
from collections import defaultdict

class ManuscriptValidator:
    def __init__(self, manuscript_path):
        self.manuscript_path = Path(manuscript_path)
        self.doc = Document(str(self.manuscript_path))
        self.validation_results = defaultdict(list)
        self.errors = []
        self.warnings = []
        self.success = []

    def get_all_text(self):
        """Extract all text from document including tables"""
        all_text = []

        # Paragraphs
        for para in self.doc.paragraphs:
            all_text.append(para.text)

        # Tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)

        return '\n'.join(all_text)

    def validate_numerical_corrections(self):
        """Validate all numerical corrections were applied"""
        print("\n" + "="*70)
        print("VALIDATION 1: NUMERICAL CORRECTIONS")
        print("="*70)

        full_text = self.get_all_text()

        # Check for CORRECT values
        correct_checks = [
            ("6.12×", "Discrimination factor"),
            ("15.3%", "Semantic discrimination"),
            ("2.5%", "Statistical discrimination"),
            ("r = 0.987", "Semantic-LLM correlation"),
            ("r = 0.988", "Statistical-LLM correlation"),
            ("Fleiss' kappa", "Kappa terminology"),
            ("κ = 0.260", "Fleiss' kappa value"),
            ("0.179", "Distinct inter-topic similarity"),
            ("0.312", "Similar inter-topic similarity"),
            ("0.358", "More Similar inter-topic similarity"),
            ("r = 0.859", "Pearson inter-rater correlation"),
            ("MAE = 0.084", "Mean Absolute Error"),
            ("+8.5%", "Grok original bias"),
            ("+2.8%", "Grok after consensus bias"),
            ("67%", "Bias reduction percentage"),
            ("17%", "Variance reduction"),
        ]

        print("\n✅ CHECKING FOR CORRECT VALUES:")
        for value, description in correct_checks:
            if value in full_text:
                self.success.append(f"  ✅ Found: {value} ({description})")
                print(f"  ✅ Found: {value} ({description})")
            else:
                self.warnings.append(f"  ⚠️  Missing: {value} ({description})")
                print(f"  ⚠️  Missing: {value} ({description})")

        # Check for INCORRECT values (should NOT exist)
        incorrect_checks = [
            ("27.3%", "Old discrimination (should be 6.12×)"),
            ("36.5%", "Old discrimination (should be 6.12×)"),
            ("r = 0.88", "Old correlation (should be 0.987)"),
            ("r = 0.67", "Old correlation (should be 0.988)"),
            ("Cohen's kappa", "Old terminology (should be Fleiss')"),
            ("Cohen's κ", "Old terminology (should be Fleiss')"),
            ("κ = 0.91", "Old kappa value (should be 0.260)"),
            ("κ = 0.89", "Old kappa value (should be 0.260)"),
        ]

        print("\n❌ CHECKING FOR INCORRECT VALUES (Should NOT exist):")
        for value, description in incorrect_checks:
            if value in full_text:
                self.errors.append(f"  ❌ FOUND OLD VALUE: {value} ({description})")
                print(f"  ❌ FOUND OLD VALUE: {value} ({description})")
            else:
                self.success.append(f"  ✅ Correctly removed: {value}")
                print(f"  ✅ Correctly removed: {value}")

    def validate_section_structure(self):
        """Validate section structure and numbering"""
        print("\n" + "="*70)
        print("VALIDATION 2: SECTION STRUCTURE")
        print("="*70)

        full_text = self.get_all_text()

        # Required sections
        required_sections = [
            ("Section 2.5", "Comparison with LLM-based Evaluation Approaches"),
            ("Section 3.1", "Experimental Data Construction"),
            ("Section 3.2.3", "Embedding Model Specification"),
            ("Section 3.3.2.1", "Parameter Configuration and Optimization"),
            ("Section 3.3.3", "LLM-based Evaluation Protocol"),
            ("Section 5.1", "Discrimination Power"),
            ("Section 5.2", "LLM Evaluation Alignment"),
            ("Section 5.3", "Methodological Limitations and Future Directions"),
            ("Section 6.1", "Key Contributions"),
            ("Section 6.2", "Limitations and Scope"),
            ("Section 6.3", "Future Research Directions"),
            ("Section 6.4", "Open Science"),
            ("Section 6.5", "Concluding Remarks"),
        ]

        print("\n📋 CHECKING REQUIRED SECTIONS:")
        for section_num, section_name in required_sections:
            # Check if section number appears
            if section_num in full_text or section_name in full_text:
                self.success.append(f"  ✅ Found: {section_num} ({section_name})")
                print(f"  ✅ Found: {section_num} ({section_name})")
            else:
                self.warnings.append(f"  ⚠️  Missing: {section_num} ({section_name})")
                print(f"  ⚠️  Missing: {section_num} ({section_name})")

    def validate_appendices(self):
        """Validate appendices existence"""
        print("\n" + "="*70)
        print("VALIDATION 3: APPENDICES")
        print("="*70)

        full_text = self.get_all_text()

        appendices = [
            ("Appendix B", "Toy Example Demonstrations"),
            ("Appendix C", "Parameter Grid Search"),
            ("Appendix D", "Wikipedia Seed Page Lists"),
            ("Appendix E", "Robustness Analysis"),
        ]

        print("\n📚 CHECKING APPENDICES:")
        for appendix_id, appendix_name in appendices:
            if appendix_id in full_text or appendix_name in full_text:
                self.success.append(f"  ✅ Found: {appendix_id} ({appendix_name})")
                print(f"  ✅ Found: {appendix_id} ({appendix_name})")
            else:
                self.warnings.append(f"  ⚠️  Missing: {appendix_id} ({appendix_name})")
                print(f"  ⚠️  Missing: {appendix_id} ({appendix_name})")

    def validate_key_content(self):
        """Validate key content additions"""
        print("\n" + "="*70)
        print("VALIDATION 4: KEY CONTENT")
        print("="*70)

        full_text = self.get_all_text()

        key_content = [
            ("October 8, 2024", "Wikipedia extraction date"),
            ("sentence-transformers/all-MiniLM-L6-v2", "Embedding model"),
            ("384", "Embedding dimensions"),
            ("γ_direct = 0.7", "Optimized parameter"),
            ("threshold_edge = 0.3", "Edge threshold"),
            ("GPT-4.1", "LLM model 1"),
            ("Claude Sonnet 4.5", "LLM model 2"),
            ("Grok", "LLM model 3"),
            ("temperature = 0.0", "LLM temperature"),
            ("3,445", "Distinct dataset size"),
            ("2,719", "Similar dataset size"),
            ("3,444", "More Similar dataset size"),
            ("142.3", "Distinct avg words"),
            ("135.8", "Similar avg words"),
            ("138.5", "More Similar avg words"),
        ]

        print("\n🔍 CHECKING KEY CONTENT:")
        for content, description in key_content:
            if content in full_text:
                self.success.append(f"  ✅ Found: {content} ({description})")
                print(f"  ✅ Found: {content} ({description})")
            else:
                self.warnings.append(f"  ⚠️  Missing: {content} ({description})")
                print(f"  ⚠️  Missing: {content} ({description})")

    def validate_references(self):
        """Validate critical cross-references"""
        print("\n" + "="*70)
        print("VALIDATION 5: CROSS-REFERENCES")
        print("="*70)

        full_text = self.get_all_text()

        references = [
            ("Appendix D", "Seed page lists reference"),
            ("Appendix C", "Grid search reference"),
            ("Appendix E", "Robustness reference"),
            ("reproducibility_guide.md", "Reproducibility guide"),
            ("Zenodo", "Data repository"),
            ("GitHub", "Code repository"),
        ]

        print("\n🔗 CHECKING CROSS-REFERENCES:")
        for ref, description in references:
            if ref in full_text:
                self.success.append(f"  ✅ Found: {ref} ({description})")
                print(f"  ✅ Found: {ref} ({description})")
            else:
                self.warnings.append(f"  ⚠️  Missing: {ref} ({description})")
                print(f"  ⚠️  Missing: {ref} ({description})")

    def generate_validation_report(self):
        """Generate comprehensive validation report"""
        print("\n" + "="*70)
        print("VALIDATION SUMMARY")
        print("="*70)

        print(f"\n✅ Successes: {len(self.success)}")
        print(f"⚠️  Warnings: {len(self.warnings)}")
        print(f"❌ Errors: {len(self.errors)}")

        # Detailed report
        report_lines = []
        report_lines.append("="*70)
        report_lines.append("MANUSCRIPT VALIDATION REPORT")
        report_lines.append("="*70)
        report_lines.append(f"\nDocument: {self.manuscript_path.name}")
        report_lines.append(f"Date: {Path(__file__).stat().st_mtime}")

        report_lines.append("\n" + "-"*70)
        report_lines.append(f"SUMMARY: {len(self.success)} ✅ | {len(self.warnings)} ⚠️ | {len(self.errors)} ❌")
        report_lines.append("-"*70)

        if self.errors:
            report_lines.append("\n❌ CRITICAL ERRORS (MUST FIX):")
            report_lines.extend(self.errors)

        if self.warnings:
            report_lines.append("\n⚠️  WARNINGS (REVIEW NEEDED):")
            report_lines.extend(self.warnings)

        if self.success:
            report_lines.append("\n✅ SUCCESSFUL VALIDATIONS:")
            report_lines.extend(self.success)

        report_lines.append("\n" + "-"*70)
        report_lines.append("NEXT STEPS:")
        report_lines.append("-"*70)

        if self.errors:
            report_lines.append("1. ❌ FIX CRITICAL ERRORS - Old values still present")
            report_lines.append("2. Review and apply missing numerical corrections")
        elif self.warnings:
            report_lines.append("1. ⚠️  REVIEW WARNINGS - Some content may be missing")
            report_lines.append("2. Manually verify missing sections/content")
            report_lines.append("3. Apply content additions from 00_MASTER_UPDATE_GUIDE.md")
        else:
            report_lines.append("1. ✅ All validations passed!")
            report_lines.append("2. Perform final manual review")
            report_lines.append("3. Prepare for journal submission")

        # Save report
        report_path = self.manuscript_path.parent / f"validation_report_{self.manuscript_path.stem}.txt"
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report_lines))

        print(f"\n📄 Validation report saved: {report_path.name}")

        # Print summary
        if self.errors:
            print("\n🚨 VALIDATION FAILED - Critical errors found!")
            print("   Please review validation_report for details")
            return False
        elif self.warnings:
            print("\n⚠️  VALIDATION INCOMPLETE - Warnings found")
            print("   Some content may need manual addition")
            print("   See 00_MASTER_UPDATE_GUIDE.md for instructions")
            return True
        else:
            print("\n✅ VALIDATION PASSED - All checks successful!")
            print("   Manuscript ready for final review")
            return True

def main():
    """Main validation execution"""
    import sys

    print("\n" + "="*70)
    print("MANUSCRIPT UPDATE VALIDATION")
    print("Phase 8: Reviewer Comment Integration")
    print("="*70)

    # Path to updated manuscript
    if len(sys.argv) > 1:
        manuscript_path = Path(sys.argv[1])
    else:
        manuscript_path = Path(r"C:\jips\docs\manuscript_updated_20251011_112640.docx")

    if not manuscript_path.exists():
        print(f"\n❌ ERROR: Updated manuscript not found: {manuscript_path}")
        print("   Please run apply_manuscript_updates.py first")
        return

    print(f"\n📄 Validating: {manuscript_path.name}")

    # Initialize validator
    validator = ManuscriptValidator(manuscript_path)

    # Run validation checks
    validator.validate_numerical_corrections()
    validator.validate_section_structure()
    validator.validate_appendices()
    validator.validate_key_content()
    validator.validate_references()

    # Generate report
    success = validator.generate_validation_report()

    print("\n" + "="*70)
    print("VALIDATION COMPLETE")
    print("="*70)

    if success:
        print("\n✅ Manuscript validation successful!")
        print("   Review validation_report for details")
    else:
        print("\n❌ Validation failed - please review errors")

if __name__ == "__main__":
    main()
