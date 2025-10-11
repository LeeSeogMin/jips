#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Phase 2 Content Application Script
Attempts to automatically apply Phase 2 content updates using python-docx

Usage: python apply_phase2_content.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from datetime import datetime

class Phase2ContentApplicator:
    def __init__(self, manuscript_path, updates_dir):
        self.manuscript_path = Path(manuscript_path)
        self.updates_dir = Path(updates_dir)
        self.doc = Document(str(self.manuscript_path))
        self.log = []
        self.changes_made = 0
        self.available_styles = [s.name for s in self.doc.styles]
        self.failed_operations = []

    def log_action(self, message, success=True):
        """Log an action with timestamp"""
        status = "✅" if success else "❌"
        log_entry = f"{status} {message}"
        self.log.append(log_entry)
        print(f"  {log_entry}")

    def find_paragraph_by_text(self, search_text, starts_with=False, contains=False):
        """Find paragraph index by text"""
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if starts_with and text.startswith(search_text):
                return i, para
            elif contains and search_text in text:
                return i, para
            elif not starts_with and not contains and text == search_text:
                return i, para
        return None, None

    def find_section_heading(self, section_number):
        """Find section heading like '3.1', '3.2', etc."""
        patterns = [
            f"{section_number} ",
            f"{section_number}.",
            f"## {section_number}",
            f"### {section_number}",
        ]

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            for pattern in patterns:
                if text.startswith(pattern):
                    return i, para
        return None, None

    def insert_paragraph_with_formatting(self, insert_after_idx, text, style_hint=None):
        """Insert a paragraph with appropriate formatting"""
        try:
            # Get the paragraph to insert after
            if insert_after_idx >= len(self.doc.paragraphs):
                # Append at end
                new_para = self.doc.add_paragraph(text)
            else:
                ref_para = self.doc.paragraphs[insert_after_idx]
                new_para = ref_para.insert_paragraph_before(text)

            # Apply formatting based on style hint
            if style_hint:
                if style_hint in self.available_styles:
                    new_para.style = style_hint
                else:
                    # Manual formatting
                    for run in new_para.runs:
                        if 'Heading 1' in style_hint or 'heading 1' in style_hint:
                            run.bold = True
                            run.font.size = Pt(16)
                        elif 'Heading 2' in style_hint or 'heading 2' in style_hint:
                            run.bold = True
                            run.font.size = Pt(14)
                        elif 'Heading 3' in style_hint or 'heading 3' in style_hint:
                            run.bold = True
                            run.font.size = Pt(12)
                        elif 'Heading 4' in style_hint or 'heading 4' in style_hint:
                            run.bold = True
                            run.font.size = Pt(11)

            return new_para
        except Exception as e:
            self.log_action(f"Failed to insert paragraph: {str(e)}", False)
            return None

    def parse_markdown_line(self, line):
        """Parse markdown line and return text and style hint"""
        line = line.strip()

        if line.startswith('#### '):
            return line[5:], 'Heading 4'
        elif line.startswith('### '):
            return line[4:], 'Heading 3'
        elif line.startswith('## '):
            return line[3:], 'Heading 2'
        elif line.startswith('# '):
            return line[2:], 'Heading 1'
        elif line.startswith('**') and line.endswith('**'):
            return line[2:-2], 'Bold'
        else:
            return line, None

    def insert_markdown_content(self, insert_after_idx, markdown_text):
        """Insert markdown content as formatted paragraphs"""
        lines = markdown_text.split('\n')
        inserted_count = 0
        current_idx = insert_after_idx

        for line in lines:
            if not line.strip():
                # Empty line - add paragraph break
                self.insert_paragraph_with_formatting(current_idx + inserted_count, "", None)
                inserted_count += 1
                continue

            text, style_hint = self.parse_markdown_line(line)
            if text:
                para = self.insert_paragraph_with_formatting(
                    current_idx + inserted_count,
                    text,
                    style_hint
                )
                if para:
                    inserted_count += 1

        return inserted_count

    def extract_replacement_text(self, update_file):
        """Extract replacement text from update markdown file"""
        filepath = self.updates_dir / update_file

        if not filepath.exists():
            self.log_action(f"Update file not found: {update_file}", False)
            return None

        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Find the replacement text section
        patterns = [
            r'## ✏️ REPLACEMENT TEXT.*?```\n(.*?)\n```',
            r'## ✏️ NEW SECTION TEXT.*?```\n(.*?)\n```',
            r'## ✏️ COMPLETE REPLACEMENT TEXT.*?```\n(.*?)\n```',
        ]

        for pattern in patterns:
            match = re.search(pattern, content, re.DOTALL)
            if match:
                return match.group(1)

        return None

    def apply_section_3_1_expansion(self):
        """Apply Section 3.1 expansion"""
        print("\n" + "="*70)
        print("ATTEMPTING SECTION 3.1 EXPANSION")
        print("="*70)

        try:
            # Find Section 3.1
            idx, para = self.find_section_heading("3.1")

            if idx is None:
                self.log_action("Section 3.1 not found", False)
                self.failed_operations.append("Section 3.1: Could not find section heading")
                return False

            self.log_action(f"Found Section 3.1 at paragraph {idx}")

            # Load replacement content
            new_content = self.extract_replacement_text("02_section_3_1_expansion.md")

            if new_content is None:
                self.log_action("Could not extract Section 3.1 replacement content", False)
                self.failed_operations.append("Section 3.1: Could not extract replacement text")
                return False

            self.log_action(f"Loaded replacement content ({len(new_content)} chars)")

            # Find end of section 3.1 (look for next section)
            end_idx = None
            for i in range(idx + 1, len(self.doc.paragraphs)):
                text = self.doc.paragraphs[i].text.strip()
                # Look for next major section
                if re.match(r'^3\.2|^4\.|^##\s*3\.2|^##\s*4\.', text):
                    end_idx = i
                    break

            if end_idx is None:
                self.log_action("Could not find end of Section 3.1, skipping", False)
                self.failed_operations.append("Section 3.1: Could not determine section boundaries")
                return False

            self.log_action(f"Section 3.1 ends at paragraph {end_idx}")

            # Delete old content (keep heading)
            paragraphs_to_delete = end_idx - idx - 1
            self.log_action(f"Deleting {paragraphs_to_delete} paragraphs")

            for i in range(paragraphs_to_delete):
                p = self.doc.paragraphs[idx + 1]._element
                p.getparent().remove(p)

            # Insert new content
            inserted = self.insert_markdown_content(idx + 1, new_content)
            self.log_action(f"Inserted {inserted} new paragraphs")

            self.changes_made += 1
            return True

        except Exception as e:
            self.log_action(f"Section 3.1 expansion failed: {str(e)}", False)
            self.failed_operations.append(f"Section 3.1: {str(e)}")
            return False

    def apply_section_3_3_additions(self):
        """Apply Section 3.3 additions (3 parts)"""
        print("\n" + "="*70)
        print("ATTEMPTING SECTION 3.3 ADDITIONS")
        print("="*70)

        try:
            filepath = self.updates_dir / "03_section_3_3_additions.md"
            if not filepath.exists():
                self.log_action("Section 3.3 update file not found", False)
                self.failed_operations.append("Section 3.3: Update file missing")
                return False

            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()

            # Extract all three parts
            parts = {
                'section_3_2_3': re.search(r'#### Section 3.2.3.*?```\n(.*?)\n```', content, re.DOTALL),
                'section_3_3_2_1': re.search(r'#### Section 3.3.2.1.*?```\n(.*?)\n```', content, re.DOTALL),
                'section_3_3_3': re.search(r'#### Section 3.3.3.*?```\n(.*?)\n```', content, re.DOTALL),
            }

            success_count = 0

            # Part A: Section 3.2.3 (NEW - insert after 3.2.2)
            if parts['section_3_2_3']:
                idx_3_2_2, _ = self.find_section_heading("3.2.2")
                if idx_3_2_2:
                    # Find end of 3.2.2
                    end_idx = None
                    for i in range(idx_3_2_2 + 1, len(self.doc.paragraphs)):
                        text = self.doc.paragraphs[i].text.strip()
                        if re.match(r'^3\.3|^3\.2\.3|^##\s*3\.3', text):
                            end_idx = i
                            break

                    if end_idx:
                        inserted = self.insert_markdown_content(
                            end_idx,
                            parts['section_3_2_3'].group(1)
                        )
                        self.log_action(f"Section 3.2.3: Inserted {inserted} paragraphs")
                        success_count += 1
                    else:
                        self.failed_operations.append("Section 3.2.3: Could not find insertion point")
                else:
                    self.failed_operations.append("Section 3.2.3: Could not find Section 3.2.2")

            # Part B: Section 3.3.2.1 (NEW - insert after 3.3.2 heading)
            if parts['section_3_3_2_1']:
                idx_3_3_2, _ = self.find_section_heading("3.3.2")
                if idx_3_3_2:
                    inserted = self.insert_markdown_content(
                        idx_3_3_2 + 1,
                        parts['section_3_3_2_1'].group(1)
                    )
                    self.log_action(f"Section 3.3.2.1: Inserted {inserted} paragraphs")
                    success_count += 1
                else:
                    self.failed_operations.append("Section 3.3.2.1: Could not find Section 3.3.2")

            # Part C: Section 3.3.3 (REPLACE content)
            if parts['section_3_3_3']:
                idx_3_3_3, _ = self.find_section_heading("3.3.3")
                if idx_3_3_3:
                    # Find end of 3.3.3
                    end_idx = None
                    for i in range(idx_3_3_3 + 1, len(self.doc.paragraphs)):
                        text = self.doc.paragraphs[i].text.strip()
                        if re.match(r'^3\.4|^4\.|^##\s*3\.4|^##\s*4\.', text):
                            end_idx = i
                            break

                    if end_idx:
                        # Delete old content
                        for i in range(end_idx - idx_3_3_3 - 1):
                            p = self.doc.paragraphs[idx_3_3_3 + 1]._element
                            p.getparent().remove(p)

                        # Insert new content
                        inserted = self.insert_markdown_content(
                            idx_3_3_3 + 1,
                            parts['section_3_3_3'].group(1)
                        )
                        self.log_action(f"Section 3.3.3: Replaced with {inserted} paragraphs")
                        success_count += 1
                    else:
                        self.failed_operations.append("Section 3.3.3: Could not find section end")
                else:
                    self.failed_operations.append("Section 3.3.3: Could not find section heading")

            if success_count > 0:
                self.changes_made += success_count
                return True
            else:
                return False

        except Exception as e:
            self.log_action(f"Section 3.3 additions failed: {str(e)}", False)
            self.failed_operations.append(f"Section 3.3: {str(e)}")
            return False

    def apply_section_2_5_new(self):
        """Apply Section 2.5 (NEW)"""
        print("\n" + "="*70)
        print("ATTEMPTING SECTION 2.5 (NEW)")
        print("="*70)

        try:
            # Load content
            new_content = self.extract_replacement_text("04_section_2_5_related_work.md")

            if new_content is None:
                self.log_action("Could not extract Section 2.5 content", False)
                self.failed_operations.append("Section 2.5: Could not extract content")
                return False

            # Find Section 2.4 end / Section 3 start
            idx_section_3, _ = self.find_section_heading("3")

            if idx_section_3 is None:
                self.log_action("Could not find Section 3 to insert before", False)
                self.failed_operations.append("Section 2.5: Could not find insertion point")
                return False

            # Insert before Section 3
            inserted = self.insert_markdown_content(idx_section_3, new_content)
            self.log_action(f"Section 2.5: Inserted {inserted} paragraphs before Section 3")

            self.changes_made += 1
            return True

        except Exception as e:
            self.log_action(f"Section 2.5 failed: {str(e)}", False)
            self.failed_operations.append(f"Section 2.5: {str(e)}")
            return False

    def apply_section_5_updates(self):
        """Apply Section 5 updates"""
        print("\n" + "="*70)
        print("ATTEMPTING SECTION 5 UPDATES")
        print("="*70)

        try:
            filepath = self.updates_dir / "05_section_5_discussion.md"
            if not filepath.exists():
                self.log_action("Section 5 update file not found", False)
                self.failed_operations.append("Section 5: Update file missing")
                return False

            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()

            # Extract parts
            parts = {
                'section_5_1': re.search(r'### 5.1.*?```\n(.*?)\n```', content, re.DOTALL),
                'section_5_2': re.search(r'### 5.2.*?```\n(.*?)\n```', content, re.DOTALL),
                'section_5_3': re.search(r'### 5.3.*?```\n(.*?)\n```', content, re.DOTALL),
            }

            success_count = 0

            # Section 5.1 (REPLACE)
            if parts['section_5_1']:
                idx, _ = self.find_section_heading("5.1")
                if idx:
                    # Find end
                    end_idx = None
                    for i in range(idx + 1, len(self.doc.paragraphs)):
                        text = self.doc.paragraphs[i].text.strip()
                        if re.match(r'^5\.2|^6\.|^##\s*5\.2|^##\s*6\.', text):
                            end_idx = i
                            break

                    if end_idx:
                        # Delete old
                        for i in range(end_idx - idx - 1):
                            p = self.doc.paragraphs[idx + 1]._element
                            p.getparent().remove(p)

                        # Insert new
                        inserted = self.insert_markdown_content(idx + 1, parts['section_5_1'].group(1))
                        self.log_action(f"Section 5.1: Replaced with {inserted} paragraphs")
                        success_count += 1
                    else:
                        self.failed_operations.append("Section 5.1: Could not find section end")
                else:
                    self.failed_operations.append("Section 5.1: Could not find section heading")

            # Section 5.2 (REPLACE or INSERT)
            if parts['section_5_2']:
                idx, _ = self.find_section_heading("5.2")
                if idx:
                    # Replace existing
                    end_idx = None
                    for i in range(idx + 1, len(self.doc.paragraphs)):
                        text = self.doc.paragraphs[i].text.strip()
                        if re.match(r'^5\.3|^6\.|^##\s*5\.3|^##\s*6\.', text):
                            end_idx = i
                            break

                    if end_idx:
                        for i in range(end_idx - idx - 1):
                            p = self.doc.paragraphs[idx + 1]._element
                            p.getparent().remove(p)

                        inserted = self.insert_markdown_content(idx + 1, parts['section_5_2'].group(1))
                        self.log_action(f"Section 5.2: Replaced with {inserted} paragraphs")
                        success_count += 1
                    else:
                        self.failed_operations.append("Section 5.2: Could not find section end")
                else:
                    # Insert new section after 5.1
                    idx_5_1, _ = self.find_section_heading("5.1")
                    if idx_5_1:
                        # Find end of 5.1
                        end_idx = None
                        for i in range(idx_5_1 + 1, len(self.doc.paragraphs)):
                            text = self.doc.paragraphs[i].text.strip()
                            if re.match(r'^5\.2|^5\.3|^6\.|^##\s*5', text):
                                end_idx = i
                                break

                        if end_idx:
                            inserted = self.insert_markdown_content(end_idx, parts['section_5_2'].group(1))
                            self.log_action(f"Section 5.2: Inserted {inserted} paragraphs")
                            success_count += 1
                        else:
                            self.failed_operations.append("Section 5.2: Could not find insertion point")
                    else:
                        self.failed_operations.append("Section 5.2: Could not find Section 5.1")

            # Section 5.3 (NEW)
            if parts['section_5_3']:
                idx_5_2, _ = self.find_section_heading("5.2")
                if idx_5_2:
                    # Find end of 5.2
                    end_idx = None
                    for i in range(idx_5_2 + 1, len(self.doc.paragraphs)):
                        text = self.doc.paragraphs[i].text.strip()
                        if re.match(r'^5\.3|^6\.|^##\s*5\.3|^##\s*6\.', text):
                            end_idx = i
                            break

                    if end_idx:
                        inserted = self.insert_markdown_content(end_idx, parts['section_5_3'].group(1))
                        self.log_action(f"Section 5.3: Inserted {inserted} paragraphs")
                        success_count += 1
                    else:
                        self.failed_operations.append("Section 5.3: Could not find insertion point")
                else:
                    self.failed_operations.append("Section 5.3: Could not find Section 5.2")

            if success_count > 0:
                self.changes_made += success_count
                return True
            else:
                return False

        except Exception as e:
            self.log_action(f"Section 5 updates failed: {str(e)}", False)
            self.failed_operations.append(f"Section 5: {str(e)}")
            return False

    def apply_section_6_updates(self):
        """Apply Section 6 updates"""
        print("\n" + "="*70)
        print("ATTEMPTING SECTION 6 UPDATES")
        print("="*70)

        try:
            # Load content
            new_content = self.extract_replacement_text("06_section_6_conclusion.md")

            if new_content is None:
                self.log_action("Could not extract Section 6 content", False)
                self.failed_operations.append("Section 6: Could not extract content")
                return False

            # Find Section 6
            idx, _ = self.find_section_heading("6")

            if idx is None:
                self.log_action("Section 6 not found", False)
                self.failed_operations.append("Section 6: Could not find section heading")
                return False

            # Find end of Section 6 (look for References or Appendix)
            end_idx = None
            for i in range(idx + 1, len(self.doc.paragraphs)):
                text = self.doc.paragraphs[i].text.strip()
                if re.match(r'^References|^Appendix|^##\s*References|^##\s*Appendix', text, re.IGNORECASE):
                    end_idx = i
                    break

            if end_idx is None:
                self.log_action("Could not find end of Section 6, skipping", False)
                self.failed_operations.append("Section 6: Could not determine section boundaries")
                return False

            # Delete old content
            paragraphs_to_delete = end_idx - idx - 1
            self.log_action(f"Deleting {paragraphs_to_delete} paragraphs from Section 6")

            for i in range(paragraphs_to_delete):
                p = self.doc.paragraphs[idx + 1]._element
                p.getparent().remove(p)

            # Insert new content
            inserted = self.insert_markdown_content(idx + 1, new_content)
            self.log_action(f"Section 6: Inserted {inserted} new paragraphs")

            self.changes_made += 1
            return True

        except Exception as e:
            self.log_action(f"Section 6 updates failed: {str(e)}", False)
            self.failed_operations.append(f"Section 6: {str(e)}")
            return False

    def apply_appendices(self):
        """Apply Appendices B, C, D, E"""
        print("\n" + "="*70)
        print("ATTEMPTING APPENDICES B, C, D, E")
        print("="*70)

        try:
            filepath = self.updates_dir / "07_appendices.md"
            if not filepath.exists():
                self.log_action("Appendices update file not found", False)
                self.failed_operations.append("Appendices: Update file missing")
                return False

            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()

            # Extract appendices
            appendices = {
                'B': re.search(r'## Appendix B:.*?```\n(.*?)\n```', content, re.DOTALL),
                'C': re.search(r'## Appendix C:.*?```\n(.*?)\n```', content, re.DOTALL),
                'D': re.search(r'## Appendix D:.*?```\n(.*?)\n```', content, re.DOTALL),
                'E': re.search(r'## Appendix E:.*?```\n(.*?)\n```', content, re.DOTALL),
            }

            # Find References section or end of document
            insert_idx = len(self.doc.paragraphs)
            for i, para in enumerate(self.doc.paragraphs):
                text = para.text.strip()
                if re.match(r'^References', text, re.IGNORECASE):
                    insert_idx = i
                    break

            success_count = 0

            # Insert appendices before References or at end
            for letter in ['B', 'C', 'D', 'E']:
                if appendices[letter]:
                    inserted = self.insert_markdown_content(
                        insert_idx,
                        appendices[letter].group(1)
                    )
                    self.log_action(f"Appendix {letter}: Inserted {inserted} paragraphs")
                    success_count += 1
                    insert_idx += inserted
                else:
                    self.log_action(f"Appendix {letter}: Content not found in update file", False)

            if success_count > 0:
                self.changes_made += success_count
                return True
            else:
                return False

        except Exception as e:
            self.log_action(f"Appendices failed: {str(e)}", False)
            self.failed_operations.append(f"Appendices: {str(e)}")
            return False

    def save_document(self):
        """Save the updated document"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.manuscript_path.parent / f"manuscript_phase2_complete_{timestamp}.docx"
        self.doc.save(str(output_path))
        self.log_action(f"Document saved: {output_path.name}")
        return output_path

    def generate_report(self):
        """Generate final report"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_path = self.manuscript_path.parent / f"phase2_application_report_{timestamp}.txt"

        report = []
        report.append("="*70)
        report.append("PHASE 2 CONTENT APPLICATION REPORT")
        report.append("="*70)
        report.append(f"\nDate: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"Source: {self.manuscript_path.name}")

        report.append("\n" + "-"*70)
        report.append(f"TOTAL CHANGES: {self.changes_made}")
        report.append("-"*70)

        report.append("\n" + "-"*70)
        report.append("APPLICATION LOG:")
        report.append("-"*70)
        for entry in self.log:
            report.append(entry)

        if self.failed_operations:
            report.append("\n" + "-"*70)
            report.append(f"FAILED OPERATIONS ({len(self.failed_operations)}):")
            report.append("-"*70)
            for failure in self.failed_operations:
                report.append(f"❌ {failure}")

            report.append("\n" + "-"*70)
            report.append("MANUAL INTERVENTION REQUIRED:")
            report.append("-"*70)
            report.append("Please refer to COMPREHENSIVE_UPDATE_GUIDE.md")
            report.append("for manual instructions on failed operations.")
        else:
            report.append("\n" + "-"*70)
            report.append("✅ ALL OPERATIONS SUCCESSFUL")
            report.append("-"*70)

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report))

        return report_path

def main():
    """Main execution"""
    print("\n" + "="*70)
    print("PHASE 2 CONTENT APPLICATION")
    print("Automated insertion of Phase 2 updates with error handling")
    print("="*70)

    manuscript_path = Path(r"C:\jips\docs\manuscript_auto_updated_20251011_114059.docx")
    updates_dir = Path(r"C:\jips\docs\manuscript_updates")

    if not manuscript_path.exists():
        print(f"\n❌ ERROR: Manuscript not found: {manuscript_path}")
        return

    applicator = Phase2ContentApplicator(manuscript_path, updates_dir)

    # Apply all updates
    operations = [
        ("Section 3.1 Expansion", applicator.apply_section_3_1_expansion),
        ("Section 3.3 Additions", applicator.apply_section_3_3_additions),
        ("Section 2.5 (NEW)", applicator.apply_section_2_5_new),
        ("Section 5 Updates", applicator.apply_section_5_updates),
        ("Section 6 Updates", applicator.apply_section_6_updates),
        ("Appendices B, C, D, E", applicator.apply_appendices),
    ]

    for name, operation in operations:
        try:
            operation()
        except Exception as e:
            print(f"\n❌ {name} failed with exception: {str(e)}")
            applicator.failed_operations.append(f"{name}: Unexpected error - {str(e)}")

    # Save document
    try:
        output_path = applicator.save_document()
        print(f"\n✅ Document saved: {output_path.name}")
    except Exception as e:
        print(f"\n❌ Failed to save document: {str(e)}")
        return

    # Generate report
    report_path = applicator.generate_report()

    print("\n" + "="*70)
    print("PHASE 2 APPLICATION COMPLETE")
    print("="*70)
    print(f"\n📄 Updated manuscript: {output_path.name}")
    print(f"📋 Application report: {report_path.name}")
    print(f"✅ Successful changes: {applicator.changes_made}")
    print(f"❌ Failed operations: {len(applicator.failed_operations)}")

    if applicator.failed_operations:
        print(f"\n⚠️  Some operations failed - manual intervention required")
        print(f"   See {report_path.name} for details")
        print(f"   Refer to COMPREHENSIVE_UPDATE_GUIDE.md for manual instructions")
    else:
        print(f"\n🎉 All operations completed successfully!")
        print(f"   Proceed to final validation")

if __name__ == "__main__":
    main()
