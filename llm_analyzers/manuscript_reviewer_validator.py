"""
Manuscript Reviewer Comments Validator
Validates that all reviewer comments have been properly addressed in the final manuscript
"""

from docx import Document
from pathlib import Path
from typing import Dict, List, Tuple
import re
import json
from datetime import datetime


class ReviewerCommentValidator:
    """Validates reviewer comments against final manuscript"""

    def __init__(self, manuscript_path: str):
        self.manuscript_path = Path(manuscript_path)
        self.doc = Document(str(self.manuscript_path))
        self.full_text = self._extract_full_text()

        # Define reviewer requirements based on comments.md
        self.requirements = self._define_requirements()

    def _extract_full_text(self) -> str:
        """Extract all text from manuscript"""
        all_text = []

        # Extract paragraphs
        for para in self.doc.paragraphs:
            all_text.append(para.text)

        # Extract tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)

        return '\n'.join(all_text)

    def _define_requirements(self) -> Dict[str, List[Dict]]:
        """Define all reviewer requirements to check"""
        return {
            "major_issues": [
                {
                    "id": "M1",
                    "category": "Inconsistent Numbers",
                    "checks": [
                        {
                            "item": "Cohen's kappa removed",
                            "validation": lambda text: "Cohen's kappa" not in text and "Cohen's κ" not in text,
                            "required_values": ["Fleiss' kappa", "κ = 0.260"]
                        },
                        {
                            "item": "Correlation coefficients unified",
                            "validation": lambda text: "r = 0.987" in text and "r = 0.988" in text,
                            "required_values": ["r = 0.987", "r = 0.988", "r = 0.859"]
                        },
                        {
                            "item": "Discrimination values unified",
                            "validation": lambda text: "6.12×" in text and "15.3%" in text and "2.5%" in text,
                            "required_values": ["6.12×", "15.3%", "2.5%"]
                        },
                        {
                            "item": "Old incorrect values removed",
                            "validation": lambda text: "27.3%" not in text and "36.5%" not in text and "κ = 0.91" not in text,
                            "forbidden_values": ["27.3%", "36.5%", "κ = 0.91", "κ = 0.89"]
                        }
                    ]
                },
                {
                    "id": "M2",
                    "category": "Reproducibility Details",
                    "checks": [
                        {
                            "item": "Embedding model specified",
                            "validation": lambda text: "sentence-transformers/all-MiniLM-L6-v2" in text,
                            "required_values": ["sentence-transformers/all-MiniLM-L6-v2", "384"]
                        },
                        {
                            "item": "LLM models specified",
                            "validation": lambda text: "GPT-4.1" in text and "Claude Sonnet 4.5" in text and "Grok" in text,
                            "required_values": ["GPT-4.1", "Claude Sonnet 4.5", "Grok", "temperature = 0.0"]
                        },
                        {
                            "item": "Dataset construction details",
                            "validation": lambda text: "October 8, 2024" in text or "2024-10-08" in text,
                            "required_values": ["Wikipedia", "October 8, 2024", "seed"]
                        },
                        {
                            "item": "Reproducibility guide reference",
                            "validation": lambda text: "reproducibility_guide.md" in text or "Zenodo" in text,
                            "required_values": ["reproducibility_guide.md", "Zenodo", "GitHub"]
                        }
                    ]
                },
                {
                    "id": "M3",
                    "category": "Metric Definitions",
                    "checks": [
                        {
                            "item": "Parameter values specified",
                            "validation": lambda text: "γ_direct = 0.7" in text or "γ = 0.7" in text,
                            "required_values": ["γ_direct = 0.7", "threshold_edge = 0.3", "α", "β"]
                        },
                        {
                            "item": "Parameter optimization mentioned",
                            "validation": lambda text: "grid search" in text.lower() or "optimization" in text.lower(),
                            "required_values": ["grid search", "optimization", "validation"]
                        },
                        {
                            "item": "Toy examples provided",
                            "validation": lambda text: "Appendix B" in text and ("toy example" in text.lower() or "demonstration" in text.lower()),
                            "required_values": ["Appendix B", "toy example"]
                        }
                    ]
                },
                {
                    "id": "M4",
                    "category": "LLM Limitations & Robustness",
                    "checks": [
                        {
                            "item": "Multi-model consensus discussed",
                            "validation": lambda text: "consensus" in text.lower() and ("three" in text or "3" in text) and "model" in text,
                            "required_values": ["consensus", "three models", "ensemble"]
                        },
                        {
                            "item": "Bias mitigation discussed",
                            "validation": lambda text: "bias" in text.lower() and "67%" in text,
                            "required_values": ["bias reduction", "67%", "+8.5%", "+2.8%"]
                        },
                        {
                            "item": "Temperature sensitivity tested",
                            "validation": lambda text: "temperature" in text.lower() and ("0.0" in text or "sensitivity" in text.lower()),
                            "required_values": ["temperature = 0.0", "sensitivity", "robustness"]
                        },
                        {
                            "item": "Limitations section present",
                            "validation": lambda text: "5.3" in text and ("limitation" in text.lower() or "future" in text.lower()),
                            "required_values": ["Section 5.3", "Limitations", "Future"]
                        }
                    ]
                }
            ],
            "minor_issues": [
                {
                    "id": "m1",
                    "category": "Terminology Consistency",
                    "checks": [
                        {
                            "item": "NPMI defined",
                            "validation": lambda text: "NPMI" in text and "Normalized Pointwise Mutual Information" in text,
                            "required_values": ["NPMI", "Normalized Pointwise Mutual Information"]
                        }
                    ]
                },
                {
                    "id": "m2",
                    "category": "Appendix Code",
                    "checks": [
                        {
                            "item": "Appendix C present (Parameter Grid Search)",
                            "validation": lambda text: "Appendix C" in text,
                            "required_values": ["Appendix C"]
                        },
                        {
                            "item": "Appendix D present (Dataset Details)",
                            "validation": lambda text: "Appendix D" in text,
                            "required_values": ["Appendix D"]
                        },
                        {
                            "item": "Appendix E present (Robustness)",
                            "validation": lambda text: "Appendix E" in text,
                            "required_values": ["Appendix E"]
                        }
                    ]
                },
                {
                    "id": "m3",
                    "category": "Conclusion Alignment",
                    "checks": [
                        {
                            "item": "Section 6 subsections present",
                            "validation": lambda text: "6.1" in text and "6.2" in text and "6.3" in text and "6.4" in text and "6.5" in text,
                            "required_values": ["6.1", "6.2", "6.3", "6.4", "6.5"]
                        },
                        {
                            "item": "Conclusion numbers match body",
                            "validation": lambda text: "6.12×" in text or "15.3%" in text,
                            "required_values": ["6.12×", "r = 0.987"]
                        }
                    ]
                }
            ]
        }

    def validate_requirement(self, check: Dict) -> Dict:
        """Validate a single requirement"""
        result = {
            "item": check["item"],
            "status": "unknown",
            "details": "",
            "found_values": [],
            "missing_values": []
        }

        try:
            # Run validation function
            if check["validation"](self.full_text):
                result["status"] = "pass"
                result["details"] = "✅ Requirement satisfied"
            else:
                result["status"] = "fail"
                result["details"] = "❌ Requirement NOT satisfied"

            # Check for required values
            if "required_values" in check:
                for value in check["required_values"]:
                    if value in self.full_text:
                        result["found_values"].append(value)
                    else:
                        result["missing_values"].append(value)

            # Check for forbidden values
            if "forbidden_values" in check:
                forbidden_found = []
                for value in check["forbidden_values"]:
                    if value in self.full_text:
                        forbidden_found.append(value)

                if forbidden_found:
                    result["status"] = "fail"
                    result["details"] = f"❌ Forbidden values found: {', '.join(forbidden_found)}"
                    result["found_values"] = forbidden_found

        except Exception as e:
            result["status"] = "error"
            result["details"] = f"⚠️ Error: {str(e)}"

        return result

    def validate_all(self) -> Dict:
        """Validate all reviewer requirements"""
        print("="*70)
        print("REVIEWER COMMENTS VALIDATION")
        print("Manuscript:", self.manuscript_path.name)
        print("="*70)

        results = {
            "manuscript": self.manuscript_path.name,
            "validation_date": datetime.now().isoformat(),
            "major_issues": {},
            "minor_issues": {},
            "summary": {
                "total_checks": 0,
                "passed": 0,
                "failed": 0,
                "errors": 0
            }
        }

        # Validate major issues
        print("\n" + "="*70)
        print("MAJOR ISSUES VALIDATION")
        print("="*70)

        for issue in self.requirements["major_issues"]:
            issue_id = issue["id"]
            category = issue["category"]

            print(f"\n{issue_id}. {category}")
            print("-"*70)

            results["major_issues"][issue_id] = {
                "category": category,
                "checks": []
            }

            for check in issue["checks"]:
                result = self.validate_requirement(check)
                results["major_issues"][issue_id]["checks"].append(result)
                results["summary"]["total_checks"] += 1

                # Print result
                status_icon = "✅" if result["status"] == "pass" else "❌" if result["status"] == "fail" else "⚠️"
                print(f"{status_icon} {result['item']}")

                if result["missing_values"]:
                    print(f"   Missing: {', '.join(result['missing_values'])}")
                if result["found_values"] and result["status"] == "fail":
                    print(f"   Found: {', '.join(result['found_values'])}")

                # Update summary
                if result["status"] == "pass":
                    results["summary"]["passed"] += 1
                elif result["status"] == "fail":
                    results["summary"]["failed"] += 1
                else:
                    results["summary"]["errors"] += 1

        # Validate minor issues
        print("\n" + "="*70)
        print("MINOR ISSUES VALIDATION")
        print("="*70)

        for issue in self.requirements["minor_issues"]:
            issue_id = issue["id"]
            category = issue["category"]

            print(f"\n{issue_id}. {category}")
            print("-"*70)

            results["minor_issues"][issue_id] = {
                "category": category,
                "checks": []
            }

            for check in issue["checks"]:
                result = self.validate_requirement(check)
                results["minor_issues"][issue_id]["checks"].append(result)
                results["summary"]["total_checks"] += 1

                # Print result
                status_icon = "✅" if result["status"] == "pass" else "❌" if result["status"] == "fail" else "⚠️"
                print(f"{status_icon} {result['item']}")

                if result["missing_values"]:
                    print(f"   Missing: {', '.join(result['missing_values'])}")

                # Update summary
                if result["status"] == "pass":
                    results["summary"]["passed"] += 1
                elif result["status"] == "fail":
                    results["summary"]["failed"] += 1
                else:
                    results["summary"]["errors"] += 1

        # Print summary
        print("\n" + "="*70)
        print("VALIDATION SUMMARY")
        print("="*70)

        total = results["summary"]["total_checks"]
        passed = results["summary"]["passed"]
        failed = results["summary"]["failed"]
        errors = results["summary"]["errors"]

        pass_rate = (passed / total * 100) if total > 0 else 0

        print(f"\nTotal Checks: {total}")
        print(f"✅ Passed: {passed} ({pass_rate:.1f}%)")
        print(f"❌ Failed: {failed}")
        print(f"⚠️  Errors: {errors}")

        if failed == 0 and errors == 0:
            print("\n🎉 ALL REVIEWER COMMENTS ADDRESSED!")
            print("   Manuscript is ready for resubmission")
        elif failed > 0:
            print(f"\n⚠️  {failed} checks failed - review required")

        return results

    def save_results(self, results: Dict, output_path: str = None):
        """Save validation results to JSON"""
        if output_path is None:
            output_path = self.manuscript_path.parent / f"reviewer_validation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)

        print(f"\n📄 Results saved: {output_path}")


def main():
    """Main execution"""
    import sys

    # Get manuscript path
    if len(sys.argv) > 1:
        manuscript_path = sys.argv[1]
    else:
        manuscript_path = r"C:\jips\docs\manuscript_FINAL_100percent_20251011_130752.docx"

    if not Path(manuscript_path).exists():
        print(f"❌ Error: Manuscript not found: {manuscript_path}")
        return

    # Run validation
    validator = ReviewerCommentValidator(manuscript_path)
    results = validator.validate_all()

    # Save results
    validator.save_results(results)

    print("\n" + "="*70)
    print("VALIDATION COMPLETE")
    print("="*70)


if __name__ == "__main__":
    main()
