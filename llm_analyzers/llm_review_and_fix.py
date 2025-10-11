#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Multi-LLM Reviewer Validation with Claude Verification and Auto-Fix
1. Run 4 LLMs to validate manuscript against reviewer comments
2. Claude verifies validity of each LLM suggestion
3. Apply only valid fixes to manuscript
"""

from docx import Document
from pathlib import Path
import json
from datetime import datetime
from typing import Dict, List, Tuple
import re

# Import LLM analyzers
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from anthropic import Anthropic
from openai import OpenAI
import google.generativeai as genai
import os
from dotenv import load_dotenv

load_dotenv()


class LLMReviewValidator:
    """Multi-LLM validation with Claude verification"""

    def __init__(self, manuscript_path: str, comments_path: str):
        self.manuscript_path = Path(manuscript_path)
        self.comments_path = Path(comments_path)
        self.doc = Document(str(self.manuscript_path))
        self.full_text = self._extract_text()

        # Read reviewer comments
        with open(self.comments_path, 'r', encoding='utf-8') as f:
            self.reviewer_comments = f.read()

        # Initialize LLMs
        self.anthropic_client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        self.openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        self.gemini_model = genai.GenerativeModel("gemini-2.5-pro")
        self.grok_client = OpenAI(api_key=os.getenv("GROK_API_KEY"), base_url="https://api.x.ai/v1")

        self.llms = {
            "Anthropic Claude": self._call_anthropic,
            "OpenAI GPT-4": self._call_openai,
            "Google Gemini": self._call_gemini,
            "xAI Grok": self._call_grok
        }

    def _call_anthropic(self, prompt: str) -> str:
        """Call Anthropic Claude API"""
        message = self.anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=4000,
            temperature=0.0,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text

    def _call_openai(self, prompt: str) -> str:
        """Call OpenAI GPT-4 API"""
        response = self.openai_client.chat.completions.create(
            model="gpt-4.1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=4000
        )
        return response.choices[0].message.content

    def _call_gemini(self, prompt: str) -> str:
        """Call Google Gemini API"""
        response = self.gemini_model.generate_content(prompt)
        return response.text

    def _call_grok(self, prompt: str) -> str:
        """Call xAI Grok API with retry logic"""
        import time

        max_retries = 3
        base_timeout = 90.0

        for attempt in range(max_retries):
            try:
                timeout = base_timeout * (attempt + 1)  # Increase timeout each retry
                print(f"  Attempt {attempt + 1}/{max_retries} (timeout: {timeout}s)...")

                response = self.grok_client.chat.completions.create(
                    model="grok-4-0709",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.0,
                    max_tokens=4000,
                    timeout=timeout
                )
                return response.choices[0].message.content

            except Exception as e:
                print(f"  ⚠️ Attempt {attempt + 1} failed: {type(e).__name__}")
                if attempt < max_retries - 1:
                    wait_time = 5 * (attempt + 1)
                    print(f"  Waiting {wait_time}s before retry...")
                    time.sleep(wait_time)
                else:
                    raise Exception(f"All {max_retries} attempts failed: {str(e)}")

    def _extract_text(self) -> str:
        """Extract all text from manuscript"""
        all_text = []
        for para in self.doc.paragraphs:
            all_text.append(para.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)
        return '\n'.join(all_text)

    def create_review_prompt(self) -> str:
        """Create prompt for LLMs to review manuscript"""

        # Use full reviewer comments
        reviewer_comments = self.reviewer_comments

        # For full manuscript, use smart sampling to stay within token limits
        # Include: Abstract, Introduction, key sections, and Conclusion
        manuscript_text = self._get_smart_manuscript_sample()

        prompt = f"""You are an expert peer reviewer. Review this manuscript against the reviewer comments below.

REVIEWER COMMENTS (COMPLETE):
{reviewer_comments}

MANUSCRIPT (COMPLETE - with key sections):
{manuscript_text}

Your task:
1. Check if ALL reviewer comments have been properly addressed in the manuscript
2. Identify any remaining issues or missing corrections
3. Provide specific, actionable fixes with exact text locations

IMPORTANT:
- The manuscript above includes ALL major sections
- Check numbers, methodology, metrics, discussions carefully
- Only report issues that are ACTUALLY missing or incorrect in the manuscript
- If a reviewer comment has been addressed, do NOT report it as an issue

Format your response as:
ISSUE_ID: [unique identifier, e.g., R1_C1]
CATEGORY: [MAJOR/MINOR]
LOCATION: [section or paragraph identifier]
CURRENT_TEXT: [exact current text that needs fixing, or state MISSING if not found]
REQUIRED_FIX: [exact text replacement or addition needed]
REASONING: [why this fix is necessary based on reviewer comments]
CONFIDENCE: [HIGH/MEDIUM/LOW]

List all issues found. If all comments addressed, state: "ALL_ISSUES_RESOLVED"
"""
        return prompt

    def _get_smart_manuscript_sample(self) -> str:
        """Return full manuscript text for complete review"""
        # Return full manuscript - LLMs can handle 60-100K characters
        # This ensures complete and accurate review
        return self.full_text

    def run_single_llm_review(self, llm_name: str) -> Dict:
        """Run single LLM review and save result"""
        print("="*70)
        print(f"RUNNING {llm_name.upper()} REVIEW")
        print("="*70)
        print(f"Manuscript: {self.manuscript_path.name}")
        print(f"Comments: {self.comments_path.name}")
        print("="*70)

        prompt = self.create_review_prompt()

        try:
            llm_func = self.llms[llm_name]
            print(f"\nCalling {llm_name} API...")
            response = llm_func(prompt)
            parsed_issues = self.parse_llm_response(response)

            result = {
                "llm_name": llm_name,
                "timestamp": datetime.now().isoformat(),
                "success": True,
                "response": response,
                "issues": parsed_issues,
                "issue_count": len(parsed_issues)
            }
            print(f"✅ {llm_name}: {len(parsed_issues)} issues found")

            # Save individual result
            self._save_individual_result(llm_name, result)

            return result

        except Exception as e:
            print(f"❌ {llm_name} failed: {e}")
            result = {
                "llm_name": llm_name,
                "timestamp": datetime.now().isoformat(),
                "success": False,
                "error": str(e),
                "issues": []
            }

            # Save error result too
            self._save_individual_result(llm_name, result)

            return result

    def _save_individual_result(self, llm_name: str, result: Dict):
        """Save individual LLM result"""
        safe_name = llm_name.replace(" ", "_").lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = self.manuscript_path.parent / f"llm_review_{safe_name}_{timestamp}.json"

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        print(f"💾 Saved: {output_path.name}")

    def run_llm_reviews(self) -> Dict[str, Dict]:
        """Run all 4 LLMs to review manuscript"""
        print("="*70)
        print("MULTI-LLM MANUSCRIPT REVIEW")
        print("="*70)
        print(f"Manuscript: {self.manuscript_path.name}")
        print(f"Comments: {self.comments_path.name}")
        print("="*70)

        results = {}

        for i, llm_name in enumerate(self.llms.keys(), start=1):
            print(f"\n[{i}/4] Running {llm_name} review...")
            result = self.run_single_llm_review(llm_name)
            results[llm_name] = result

        return results

    def parse_llm_response(self, response: str) -> List[Dict]:
        """Parse structured LLM response into issues"""
        if "ALL_ISSUES_RESOLVED" in response:
            return []

        issues = []
        current_issue = {}

        for line in response.split('\n'):
            line = line.strip()

            if line.startswith("ISSUE_ID:"):
                if current_issue:
                    issues.append(current_issue)
                current_issue = {"issue_id": line.split(":", 1)[1].strip()}
            elif line.startswith("CATEGORY:"):
                current_issue["category"] = line.split(":", 1)[1].strip()
            elif line.startswith("LOCATION:"):
                current_issue["location"] = line.split(":", 1)[1].strip()
            elif line.startswith("CURRENT_TEXT:"):
                current_issue["current_text"] = line.split(":", 1)[1].strip()
            elif line.startswith("REQUIRED_FIX:"):
                current_issue["required_fix"] = line.split(":", 1)[1].strip()
            elif line.startswith("REASONING:"):
                current_issue["reasoning"] = line.split(":", 1)[1].strip()
            elif line.startswith("CONFIDENCE:"):
                current_issue["confidence"] = line.split(":", 1)[1].strip()

        if current_issue:
            issues.append(current_issue)

        return issues

    def verify_issue_with_claude(self, issue: Dict, llm_name: str) -> Dict:
        """Use Claude to verify if suggested fix is valid and necessary"""
        verification_prompt = f"""You are Claude, an expert reviewer verifying another LLM's manuscript review suggestion.

LLM: {llm_name}
SUGGESTED ISSUE:
- ID: {issue.get('issue_id', 'N/A')}
- Category: {issue.get('category', 'N/A')}
- Location: {issue.get('location', 'N/A')}
- Current Text: {issue.get('current_text', 'N/A')}
- Required Fix: {issue.get('required_fix', 'N/A')}
- Reasoning: {issue.get('reasoning', 'N/A')}
- Confidence: {issue.get('confidence', 'N/A')}

REVIEWER COMMENTS EXCERPT:
{self.reviewer_comments[:1500]}

MANUSCRIPT EXCERPT (around suggested location):
{self._get_context_around_issue(issue)}

Your task: Verify if this suggested fix is:
1. VALID: Does it correctly address a reviewer comment?
2. NECESSARY: Is the current text actually problematic?
3. ACCURATE: Is the suggested fix correct and complete?
4. SAFE: Will this fix not introduce new errors?

Respond with:
VERDICT: [VALID/INVALID/UNCERTAIN]
VALIDITY_SCORE: [0-10, where 10=definitely valid]
REASONING: [your detailed reasoning]
RECOMMENDATION: [APPLY/REJECT/MODIFY]
MODIFIED_FIX: [if MODIFY, provide corrected fix; otherwise copy original]
"""

        try:
            response = self._call_anthropic(verification_prompt)
            verification = self.parse_verification(response)
            verification["llm_source"] = llm_name
            verification["original_issue"] = issue
            return verification
        except Exception as e:
            return {
                "verdict": "ERROR",
                "validity_score": 0,
                "reasoning": f"Verification error: {str(e)}",
                "recommendation": "REJECT",
                "llm_source": llm_name,
                "original_issue": issue
            }

    def _get_context_around_issue(self, issue: Dict) -> str:
        """Get manuscript context around the issue location"""
        current_text = issue.get('current_text', '')
        if not current_text:
            return self.full_text[:1000]

        # Find the text in manuscript
        idx = self.full_text.find(current_text[:100])
        if idx == -1:
            return self.full_text[:1000]

        # Return context: 500 chars before and after
        start = max(0, idx - 500)
        end = min(len(self.full_text), idx + len(current_text) + 500)
        return self.full_text[start:end]

    def parse_verification(self, response: str) -> Dict:
        """Parse Claude's verification response"""
        verification = {}

        for line in response.split('\n'):
            line = line.strip()
            if line.startswith("VERDICT:"):
                verification["verdict"] = line.split(":", 1)[1].strip()
            elif line.startswith("VALIDITY_SCORE:"):
                try:
                    verification["validity_score"] = float(line.split(":", 1)[1].strip())
                except:
                    verification["validity_score"] = 0.0
            elif line.startswith("REASONING:"):
                verification["reasoning"] = line.split(":", 1)[1].strip()
            elif line.startswith("RECOMMENDATION:"):
                verification["recommendation"] = line.split(":", 1)[1].strip()
            elif line.startswith("MODIFIED_FIX:"):
                verification["modified_fix"] = line.split(":", 1)[1].strip()

        return verification

    def aggregate_and_verify_issues(self, llm_results: Dict) -> List[Dict]:
        """Aggregate issues from all LLMs and verify with Claude"""
        print("\n" + "="*70)
        print("CLAUDE VERIFICATION OF LLM SUGGESTIONS")
        print("="*70)

        all_issues = []

        for llm_name, result in llm_results.items():
            if not result.get("success"):
                continue

            issues = result.get("issues", [])
            print(f"\n{llm_name}: {len(issues)} issues")

            for issue in issues:
                print(f"  Verifying: {issue.get('issue_id', 'N/A')}...", end=" ")
                verification = self.verify_issue_with_claude(issue, llm_name)

                verdict = verification.get("verdict", "UNCERTAIN")
                score = verification.get("validity_score", 0)
                recommendation = verification.get("recommendation", "REJECT")

                # Add verification to issue
                issue["verification"] = verification

                # Only include if Claude recommends APPLY or MODIFY
                if recommendation in ["APPLY", "MODIFY"]:
                    all_issues.append(issue)
                    print(f"✅ {verdict} (score: {score:.1f}) - {recommendation}")
                else:
                    print(f"❌ {verdict} (score: {score:.1f}) - REJECTED")

        return all_issues

    def apply_valid_fixes(self, verified_issues: List[Dict]) -> Tuple[int, List[str]]:
        """Apply only verified fixes to manuscript"""
        print("\n" + "="*70)
        print("APPLYING VERIFIED FIXES")
        print("="*70)

        applied_count = 0
        applied_fixes = []

        for issue in verified_issues:
            verification = issue.get("verification", {})
            recommendation = verification.get("recommendation", "REJECT")

            if recommendation not in ["APPLY", "MODIFY"]:
                continue

            current_text = issue.get("current_text", "")

            if recommendation == "MODIFY":
                required_fix = verification.get("modified_fix", issue.get("required_fix", ""))
            else:
                required_fix = issue.get("required_fix", "")

            if not current_text or not required_fix:
                continue

            # Apply fix
            success = self._apply_fix_to_doc(current_text, required_fix)

            if success:
                applied_count += 1
                applied_fixes.append({
                    "issue_id": issue.get("issue_id", "N/A"),
                    "category": issue.get("category", "N/A"),
                    "current": current_text[:100],
                    "fixed": required_fix[:100],
                    "llm_source": verification.get("llm_source", "N/A")
                })
                print(f"✅ Applied: {issue.get('issue_id', 'N/A')}")
            else:
                print(f"⚠️  Failed to apply: {issue.get('issue_id', 'N/A')}")

        return applied_count, applied_fixes

    def _apply_fix_to_doc(self, current_text: str, required_fix: str) -> bool:
        """Apply fix to document"""
        try:
            # Find paragraph containing current text
            for para in self.doc.paragraphs:
                if current_text[:50] in para.text:
                    old_text = para.text
                    new_text = old_text.replace(current_text, required_fix)
                    if new_text != old_text:
                        para.text = new_text
                        return True
            return False
        except Exception as e:
            print(f"      Error: {e}")
            return False

    def save_results(self, llm_results: Dict, verified_issues: List[Dict],
                    applied_fixes: List[Dict], output_manuscript_path: Path):
        """Save all results"""
        # Save fixed manuscript
        self.doc.save(str(output_manuscript_path))
        print(f"\n✅ Fixed manuscript saved: {output_manuscript_path.name}")

        # Save JSON results
        results = {
            "timestamp": datetime.now().isoformat(),
            "original_manuscript": self.manuscript_path.name,
            "fixed_manuscript": output_manuscript_path.name,
            "llm_reviews": {
                name: {
                    "success": result.get("success", False),
                    "issue_count": result.get("issue_count", 0)
                }
                for name, result in llm_results.items()
            },
            "verified_issues": len(verified_issues),
            "applied_fixes": len(applied_fixes),
            "fix_details": applied_fixes
        }

        json_path = output_manuscript_path.parent / f"llm_review_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)

        print(f"✅ Results saved: {json_path.name}")

    def run_full_validation(self):
        """Run complete validation and fixing workflow"""
        print("\n" + "="*70)
        print("STARTING MULTI-LLM REVIEW WITH CLAUDE VERIFICATION")
        print("="*70)

        # Step 1: Run 4 LLM reviews
        llm_results = self.run_llm_reviews()

        # Step 2: Verify with Claude
        verified_issues = self.aggregate_and_verify_issues(llm_results)

        print(f"\n📊 Summary:")
        total_issues = sum(r.get("issue_count", 0) for r in llm_results.values() if r.get("success"))
        print(f"   Total issues found: {total_issues}")
        print(f"   Verified as valid: {len(verified_issues)}")
        print(f"   Verification rate: {len(verified_issues)/total_issues*100:.1f}%" if total_issues > 0 else "   Verification rate: N/A")

        # Step 3: Apply verified fixes
        if verified_issues:
            applied_count, applied_fixes = self.apply_valid_fixes(verified_issues)
            print(f"\n✅ Applied {applied_count} fixes to manuscript")

            # Step 4: Save results
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.manuscript_path.parent / f"manuscript_FIXED_{timestamp}.docx"
            self.save_results(llm_results, verified_issues, applied_fixes, output_path)

            print("\n" + "="*70)
            print("REVIEW AND FIXING COMPLETE")
            print("="*70)
            print(f"\n📄 Fixed manuscript: {output_path.name}")
            print(f"📊 Applied fixes: {applied_count}")

            return output_path
        else:
            print("\n🎉 No valid fixes needed - manuscript is ready!")
            return None


def main():
    """Main execution"""
    import sys

    if len(sys.argv) > 2:
        manuscript_path = sys.argv[1]
        comments_path = sys.argv[2]
    else:
        manuscript_path = "docs/manuscript_F.docx"
        comments_path = "docs/comments.md"

    if not Path(manuscript_path).exists():
        print(f"❌ Error: Manuscript not found: {manuscript_path}")
        return

    if not Path(comments_path).exists():
        print(f"❌ Error: Comments not found: {comments_path}")
        return

    validator = LLMReviewValidator(manuscript_path, comments_path)
    validator.run_full_validation()


if __name__ == "__main__":
    main()
